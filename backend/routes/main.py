jus# Main API Routes
import json
import requests
import time
import os
from fastapi import FastAPI, HTTPException, UploadFile, File, Query, Depends, status
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel
from typing import List

from models import get_db, ChatHistory, UploadedDocument, UploadedCSV
from auth import get_current_user
from typing import Optional
from models import User
from auth import get_current_user
from config import ALLOWED_OLLAMA_MODELS, API_KEYS, BASE_URL
from data_analysis import (
    query_ollama_data_analysis, generate_fallback_code, execute_data_analysis_code,
    save_to_sql_database, load_csv_file, get_database_engine
)
from sqlalchemy import inspect

# Pydantic Models
class PromptInput(BaseModel):
    prompt: str
    model: str
    session_id: str

class ChainRequest(BaseModel):
    prompt: str
    models: list[str]

class DataAnalysisRequest(BaseModel):
    prompt: str
    session_id: str
    model: str = "deepseek-coder-v2:latest"

class CSVUploadResponse(BaseModel):
    filename: str
    shape: tuple[int, int]
    columns: list[str]
    dtypes: dict[str, str]
    sample_data: list[dict]

def safe_string_conversion(value, default=""):
    """Safely convert any value to string, handling lists and other types"""
    if isinstance(value, str):
        return value
    elif isinstance(value, list):
        return ' '.join(str(item) for item in value if item)
    elif value is None:
        return default
    else:
        return str(value)

def register_main_routes(app: FastAPI):
    
    @app.get("/health")
    def health_check():
        return {"status": "healthy", "message": "Backend is running"}

    @app.get("/debug/keys")
    def debug_api_keys():
        """Diagnostic endpoint to check API key parsing"""
        raw_keys_str = os.getenv("OPENROUTER_API_KEYS", "")
        parsed_keys = [key.strip() for key in raw_keys_str.split(",") if key.strip()]
        
        return {
            "raw_environment_variable_length": len(raw_keys_str),
            "raw_environment_variable_preview": raw_keys_str[:100] + "..." if len(raw_keys_str) > 100 else raw_keys_str,
            "total_keys_after_parsing": len(parsed_keys),
            "parsed_keys_preview": [key[:10] + "..." + key[-4:] if len(key) > 14 else key for key in parsed_keys[:5]],
            "all_parsed_keys_count": len(parsed_keys),
            "environment_variable_name": "OPENROUTER_API_KEYS"
        }

@app.post("/ask")
    def ask_model(data: PromptInput, db=Depends(get_db), current_user: Optional[User] = Depends(get_current_user)):
    print(f"Ask request from anon/bot session {data.session_id}")
    
    # Fetch history/docs only for authenticated users
    if current_user is not None:
        # Fetch last 1 message for context
        history = db.query(ChatHistory)\
            .filter(ChatHistory.user_id == current_user.id, ChatHistory.session_id == data.session_id)\
            .order_by(ChatHistory.timestamp.desc())\
            .limit(1).all()
    else:
        history = []

        context = ""
        for msg in reversed(history):
            context += f"User: {msg.prompt}\nAI: {msg.response}\n"

    # Fetch PDF documents (only for auth users)
    if current_user is not None:
        docs = db.query(UploadedDocument).filter(
            UploadedDocument.user_id == current_user.id, 
            UploadedDocument.session_id == data.session_id, 
            UploadedDocument.extracted_text != None
        ).order_by(UploadedDocument.timestamp.asc()).all()
    else:
        docs = []

        pdf_context = ""
        for idx, doc in enumerate(docs):
            if doc.extracted_text:
                pdf_context += f"[PDF {idx+1} Content Start: {doc.filename}]\n{doc.extracted_text}\n[PDF {idx+1} Content End]\n"

        # Search SearxNG
        searxng_context = ""
        try:
            searxng_url = "https://etsi.me/search"
            params = {
                "q": data.prompt,
                "format": "json",
                "categories": "general",
                "engines": "google,bing,duckduckgo",
                "pageno": 1
            }
            
            search_response = requests.get(searxng_url, params=params, timeout=10)
            if search_response.ok:
                search_data = search_response.json()
                search_results = search_data.get("results", [])[:5]
                
                if search_results:
                    searxng_context = "\n[SEARCH RESULTS START]\n"
                    for idx, result in enumerate(search_results, 1):
                        title = result.get("title", "No title")
                        content = result.get("content", "No content")
                        url = result.get("url", "No URL")
                        searxng_context += f"Source {idx}: {title}\nURL: {url}\nContent: {content}\n\n"
                    searxng_context += "[SEARCH RESULTS END]\n\n"
        except Exception as e:
            print(f"SearxNG search failed: {e}")

        # Combine contexts
        full_prompt = ""
        if pdf_context:
            full_prompt += pdf_context
        if searxng_context:
            full_prompt += searxng_context
        
        full_prompt += context + f"Based on the above search results and context, please answer the following question: {data.prompt}\nAI:"

        # Use Ollama
        ollama_url = "http://localhost:11434/api/generate"
        payload = {
            "model": data.model,
            "prompt": full_prompt,
            "stream": True,
            "num_predict": 256,
            "temperature": 0.7,
        }

        if data.model not in ALLOWED_OLLAMA_MODELS:
            raise HTTPException(status_code=400, detail=f"Model '{data.model}' is not available on this server.")

        def stream_ollama():
            if searxng_context:
                yield json.dumps({
                    "type": "search_results", 
                    "content": searxng_context,
                    "done": False
                }) + "\n"
            
            try:
                with requests.post(ollama_url, json=payload, stream=True, timeout=600) as response:
                    response.raise_for_status()
                    for line in response.iter_lines():
                        if line:
                            try:
                                chunk = line.decode()
                                data_json = json.loads(chunk)
                                resp = data_json.get("response", "")
                                done = data_json.get("done", False)
                                yield json.dumps({
                                    "type": "model_response",
                                    "response": resp, 
                                    "done": done
                                }) + "\n"
                                if done:
                                    break
                            except Exception:
                                continue
                yield json.dumps({"type": "model_response", "done": True}) + "\n"
            except Exception as e:
                yield json.dumps({
                    "type": "error",
                    "error": f"Ollama error: {str(e)}", 
                    "done": True
                }) + "\n"

        return StreamingResponse(stream_ollama(), media_type="application/jsonl")

    @app.post("/chain")
    def chain_models(data: ChainRequest):
        def stream_chain():
            current_prompt = data.prompt
            for model_id in data.models:
                if model_id not in ALLOWED_OLLAMA_MODELS:
                    yield json.dumps({"model": model_id, "response": f"Model '{model_id}' is not available on this server.", "done": True}) + "\n"
                    current_prompt = data.prompt
                    continue
                ollama_url = "http://localhost:11434/api/generate"
                payload = {
                    "model": model_id,
                    "prompt": current_prompt,
                    "stream": True,
                    "num_predict": 64,
                    "temperature": 0.8,
                }
                full_response = ""
                try:
                    with requests.post(ollama_url, json=payload, stream=True, timeout=600) as response:
                        response.raise_for_status()
                        for line in response.iter_lines():
                            if line:
                                try:
                                    chunk = line.decode()
                                    data_json = json.loads(chunk)
                                    resp = data_json.get("response", "")
                                    done = data_json.get("done", False)
                                    full_response += resp
                                    yield json.dumps({"model": model_id, "response": resp, "done": False}) + "\n"
                                    if done:
                                        break
                                except Exception:
                                    continue
                        yield json.dumps({"model": model_id, "done": True}) + "\n"
                        current_prompt = full_response
                except Exception as e:
                    yield json.dumps({"model": model_id, "response": f"Ollama error: {str(e)}", "done": True}) + "\n"
                    current_prompt = data.prompt
        return StreamingResponse(stream_chain(), media_type="application/jsonl")

    @app.post("/upload")
    def upload_files(files: List[UploadFile] = File(...), session_id: str = Query(...), db=Depends(get_db), current_user=Depends(get_current_user)):
        import shutil
        from config import UPLOAD_DIR
        from PIL import Image
        import pytesseract
        import PyPDF2
        import docx
        
        saved_files = []
        for file in files:
            if file.filename is None:
                continue
            filename = file.filename
            file_path = os.path.join(UPLOAD_DIR, filename)
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            extracted_text = None
            
            # Handle file types
            if filename.lower().endswith(".pdf"):
                try:
                    with open(file_path, "rb") as f:
                        reader = PyPDF2.PdfReader(f)
                        text = ""
                        for page in reader.pages:
                            text += page.extract_text() or ""
                        extracted_text = text
                except Exception as e:
                    extracted_text = None
            elif filename.lower().endswith(".docx"):
                try:
                    docx_file = docx.Document(file_path)
                    text = "\n".join([para.text for para in docx_file.paragraphs])
                    extracted_text = text
                except Exception as e:
                    extracted_text = None
            elif filename.lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".gif")):
                try:
                    image = Image.open(file_path)
                    text = pytesseract.image_to_string(image)
                    extracted_text = text
                except Exception as e:
                    extracted_text = None
            
            doc = UploadedDocument(
                user_id=current_user.id,
                session_id=session_id,
                filename=filename,
                extracted_text=extracted_text
            )
            db.add(doc)
            db.commit()
            db.refresh(doc)
            saved_files.append({"filename": filename, "extracted_text": bool(extracted_text)})
        return {"files": saved_files}