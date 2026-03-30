import subprocess
import os
import sys
import time
import requests
import json
import urllib.parse
import warnings
from pathlib import Path
from fastapi import Query, Depends
from contextlib import redirect_stdout
import traceback

# Defer heavy imports - load on demand
_heavy_imports_loaded = False

def ensure_heavy_imports():
    """Lazy load heavy data science libraries only when needed"""
    global _heavy_imports_loaded
    if _heavy_imports_loaded:
        return
    
    global pd, np, matplotlib, plt, sns, px, go, sqlite3, io, base64
    global LinearRegression, train_test_split, r2_score, mean_squared_error, stats
    
    import pandas as pd
    import numpy as np
    import matplotlib
    matplotlib.use('Agg')  # Use non-interactive backend
    import matplotlib.pyplot as plt
    import seaborn as sns
    import plotly.express as px
    import plotly.graph_objects as go
    import sqlite3
    import io
    import base64
    from sklearn.linear_model import LinearRegression
    from sklearn.model_selection import train_test_split
    from sklearn.metrics import r2_score, mean_squared_error
    from scipy import stats
    
    _heavy_imports_loaded = True

from sqlalchemy import create_engine, inspect


# Suppress warnings including bcrypt version warnings
warnings.filterwarnings('ignore')
import logging
logging.getLogger('passlib').setLevel(logging.ERROR)

# Defer Ollama initialization - will be done lazily on first use
_ollama_initialized = False

def ensure_ollama_running():
    """Lazy initialization of Ollama - only called when actually needed"""
    global _ollama_initialized
    if _ollama_initialized:
        return
    
    import socket
    s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    try:
        s.connect(('localhost', 11434))
        s.close()
        _ollama_initialized = True
        return
    except Exception:
        pass
    
    # Start Ollama server in the background
    try:
        subprocess.Popen(['ollama', 'serve'], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        _ollama_initialized = True
    except Exception as e:
        print(f"Warning: Could not start Ollama: {e}")

from fastapi import FastAPI, HTTPException, UploadFile, File, Query, Depends, status, Security
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
import requests
import random
import io
import os
from dotenv import load_dotenv
load_dotenv()

try:
    from ddgs import DDGS
except ImportError:
    DDGS = None

# Defer file processing imports - load on demand
_file_processing_imports_loaded = False

# Import PIL for type hints (actual lazy loading happens in ensure_file_processing_imports)
from typing import TYPE_CHECKING
if TYPE_CHECKING:
    from PIL import Image

# Supported OCR languages for Tesseract (including all Indian languages)
SUPPORTED_OCR_LANGUAGES = {
    # Indian Languages
    "eng": "English",
    "hin": "Hindi (हिन्दी)",
    "ben": "Bengali (বাংলা)",
    "tam": "Tamil (தமிழ்)",
    "tel": "Telugu (తెలుగు)",
    "mar": "Marathi (मराठी)",
    "guj": "Gujarati (ગુજરાતી)",
    "kan": "Kannada (ಕನ್ನಡ)",
    "mal": "Malayalam (മലയാളം)",
    "pan": "Punjabi (ਪੰਜਾਬੀ)",
    "ori": "Oriya (ଓଡିଆ)",
    "asm": "Assamese (অসমীয়া)",
    "urd": "Urdu (اردو)",
    "nep": "Nepali (नेपाली)",
    "san": "Sanskrit (संस्कृत)",
    "sin": "Sinhala (සිංහල)",
    "kok": "Konkani (कोंकणी)",
    "mni": "Manipuri (মৈতৈলোক)",
    "doi": "Dogri (डोगरी)",
    "snd": "Sindhi (सिंधी)",
    "sat": "Santali (ᱥᱟᱱᱛᱟᱲᱤ)",
    "ks": "Kashmiri (کٲشُر)",
    # Other Languages
    "spa": "Spanish",
    "fra": "French",
    "deu": "German",
    "chi_sim": "Chinese (Simplified)",
    "chi_tra": "Chinese (Traditional)",
    "jpn": "Japanese",
    "kor": "Korean",
    "ara": "Arabic",
    "por": "Portuguese",
    "ita": "Italian",
    "rus": "Russian",
    "nld": "Dutch",
    "pol": "Polish",
    "tur": "Turkish",
    "vie": "Vietnamese",
    "tha": "Thai",
    "ind": "Indonesian",
    "msa": "Malay"
}

# Pydantic model for OCR request
class OCRRequest(BaseModel):
    languages: str = "hin"  # Default to Hindi for better Hindi OCR accuracy
    enhance: bool = True  # Enable image preprocessing
    output_format: str = "text"  # text, json, or both

def ensure_file_processing_imports():
    """Lazy load file processing libraries only when uploading files"""
    global _file_processing_imports_loaded
    if _file_processing_imports_loaded:
        return
    
    global Image, pytesseract, PyPDF2, docx, cv2, np
    
    from PIL import Image
    import pytesseract
    import PyPDF2
    import docx
    
    # Try to import OpenCV for image preprocessing (optional)
    try:
        import cv2
        import numpy as np
    except ImportError:
        print("Warning: OpenCV not installed. Image preprocessing for OCR will be limited.")
        cv2 = None
        np = None
    
    _file_processing_imports_loaded = True

def preprocess_image_for_ocr(image_path: str, enhance: bool = True) -> "Image.Image":
    """Preprocess image for better OCR results"""
    from PIL import Image, ImageEnhance, ImageFilter
    import numpy as np
    
    img = Image.open(image_path)
    
    # Convert to grayscale if needed
    if img.mode != 'L':
        img = img.convert('L')
    
    if enhance:
        # Increase contrast
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(1.5)
        
        # Sharpen
        img = img.filter(ImageFilter.SHARPEN)
        
        # Apply thresholding using numpy
        img_array = np.array(img)
        # Otsu's thresholding
        img_array = (img_array > np.mean(img_array)).astype(np.uint8) * 255
        img = Image.fromarray(img_array)
    
    return img

def perform_multilingual_ocr(image_path: str, languages: str = "hin", enhance: bool = True) -> dict:
    """Perform OCR with multilingual support"""
    from PIL import Image
    import pytesseract
    import numpy as np
    
    result = {
        "text": "",
        "confidence": 0,
        "languages_used": languages,
        "success": False,
        "error": None
    }
    
    try:
        # Open and preprocess image
        img = preprocess_image_for_ocr(image_path, enhance=enhance)
        
        # Get detailed OCR data with confidence
        data = pytesseract.image_to_data(img, lang=languages, output_type=pytesseract.Output.DICT)
        
        # Extract text
        result["text"] = pytesseract.image_to_string(img, lang=languages)
        
        # Calculate average confidence
        confidences = [c for c in data['conf'] if c != -1]
        result["confidence"] = sum(confidences) / len(confidences) if confidences else 0
        
        result["success"] = True
        
    except Exception as e:
        result["error"] = str(e)
    
    return result

# For file processing
from typing import List
from fastapi.responses import StreamingResponse
from fastapi.responses import JSONResponse
import json
from fastapi.responses import FileResponse

from sqlalchemy import create_engine, Column, Integer, String, ForeignKey, Text, DateTime
from sqlalchemy.orm import sessionmaker, declarative_base, relationship
try:
    from passlib.hash import pbkdf2_sha256, bcrypt
except Exception:
    pbkdf2_sha256 = None
    bcrypt = None
import hashlib
import binascii
import hmac
import os
from jose import JWTError, jwt
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
import datetime
import shutil

# Determine base directory for data storage
if os.environ.get("DESKTOP_MODE") == "1":
    # In desktop mode, use the current working directory (which is set to user data dir)
    DATA_BASE_DIR = os.getcwd()
else:
    # In dev/server mode, use the directory of this file
    DATA_BASE_DIR = os.path.dirname(__file__)

UPLOAD_DIR = os.path.join(DATA_BASE_DIR, 'uploads')
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Create CSV upload directory
CSV_UPLOAD_DIR = os.path.join(DATA_BASE_DIR, 'csv_uploads')
os.makedirs(CSV_UPLOAD_DIR, exist_ok=True)

# Web Search Function with Multiple Fallbacks
def perform_web_search(query: str, max_results: int = 5):
    """Perform web search using multiple fallback methods."""
    results = []
    
    # Method 1: Try DuckDuckGo directly (most reliable)
    if DDGS:
        try:
            with DDGS() as ddgs:
                ddg_results = list(ddgs.text(query, max_results=max_results))
                for result in ddg_results:
                    results.append({
                        "title": result.get("title", ""),
                        "content": result.get("body", ""),
                        "url": result.get("href", ""),
                        "engine": "duckduckgo"
                    })
                if results:
                    return {"success": True, "results": results, "source": "duckduckgo"}
        except Exception as e:
            print(f"DuckDuckGo search failed: {e}")
    
    # Method 2: Try public SearXNG instances
    searxng_instances = [
        "https://search.bus-hit.me",
        "https://searx.be",
        "https://search.sapti.me",
        "https://paulgo.io"
    ]
    
    for instance in searxng_instances:
        try:
            params = {
                "q": query,
                "format": "json",
                "categories": "general",
                "pageno": 1
            }
            response = requests.get(f"{instance}/search", params=params, timeout=8)
            if response.ok:
                data = response.json()
                search_results = data.get("results", [])[:max_results]
                if search_results:
                    for result in search_results:
                        results.append({
                            "title": result.get("title", ""),
                            "content": result.get("content", ""),
                            "url": result.get("url", ""),
                            "engine": result.get("engine", "searxng")
                        })
                    if results:
                        return {"success": True, "results": results, "source": instance}
        except Exception as e:
            print(f"SearXNG instance {instance} failed: {e}")
            continue
    
    # Method 3: Fallback to DuckDuckGo HTML scraping (last resort)
    try:
        ddg_url = "https://html.duckduckgo.com/html/"
        response = requests.post(ddg_url, data={"q": query}, timeout=10)
        if response.ok:
            # Simple extraction - look for result links
            import re
            urls = re.findall(r'uddg=([^"&]+)', response.text)
            titles = re.findall(r'<a class="result__a"[^>]*>([^<]+)</a>', response.text)
            
            for i, (url, title) in enumerate(zip(urls[:max_results], titles[:max_results])):
                results.append({
                    "title": title.strip(),
                    "content": "",
                    "url": urllib.parse.unquote(url),
                    "engine": "duckduckgo-html"
                })
            
            if results:
                return {"success": True, "results": results, "source": "duckduckgo-html"}
    except Exception as e:
        print(f"DuckDuckGo HTML fallback failed: {e}")
    
    return {"success": False, "results": [], "error": "All search methods failed"}

# --- User Auth & RBAC Setup ---
# Use absolute path to ensure database is always in backend directory
DB_PATH = os.path.join(os.path.dirname(__file__), 'users.db')
DATABASE_URL = f"sqlite:///{DB_PATH}"
Base = declarative_base()
engine = create_engine(DATABASE_URL, connect_args={"check_same_thread": False})
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

class User(Base):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True, index=True)
    username = Column(String, unique=True, index=True, nullable=False)
    password_hash = Column(String, nullable=False)
    role = Column(String, default="user")
    web_search_enabled = Column(Integer, default=1)  # 1=enabled, 0=disabled
    created_at = Column(DateTime, default=datetime.datetime.utcnow)

class ChatHistory(Base):
    __tablename__ = "chat_history"
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    session_id = Column(String, nullable=False, index=True)
    prompt = Column(Text, nullable=False)
    response = Column(Text, nullable=False)
    model_used = Column(String, nullable=True)  # Track which model was used for this response
    timestamp = Column(DateTime, default=datetime.datetime.utcnow)
    user = relationship("User", backref="user_chat_history")

# New table for uploaded documents
class UploadedDocument(Base):
    __tablename__ = "uploaded_documents"
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    session_id = Column(String, nullable=False, index=True)
    filename = Column(String, nullable=False)
    extracted_text = Column(Text, nullable=True)
    timestamp = Column(DateTime, default=datetime.datetime.utcnow)
    user = relationship("User", backref="user_uploaded_documents")

# New table for CSV files in data analysis mode
class UploadedCSV(Base):
    __tablename__ = "uploaded_csvs"
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    session_id = Column(String, nullable=False, index=True)
    filename = Column(String, nullable=False)
    file_path = Column(String, nullable=False)
    columns_info = Column(Text, nullable=True)  # JSON string of column info
    table_name = Column(String, nullable=True)  # Name of the table where data is stored
    timestamp = Column(DateTime, default=datetime.datetime.utcnow)
    user = relationship("User", backref="user_uploaded_csvs")

# AI Models managed by admin
class AIModel(Base):
    __tablename__ = "ai_models"
    id = Column(Integer, primary_key=True, index=True)
    name = Column(String, unique=True, nullable=False)
    description = Column(Text, nullable=True)
    usecases = Column(Text, nullable=True)
    is_enabled = Column(Integer, default=1)  # 1=enabled, 0=disabled
    is_default = Column(Integer, default=0)  # 1=default model, 0=not default
    created_by = Column(Integer, ForeignKey("users.id"), nullable=True)
    created_at = Column(DateTime, default=datetime.datetime.utcnow)
    creator = relationship("User", backref="created_models")

# Server Settings for global configuration
class ServerSettings(Base):
    __tablename__ = "server_settings"
    id = Column(Integer, primary_key=True, index=True)
    setting_key = Column(String, unique=True, nullable=False)
    setting_value = Column(String, nullable=False)
    updated_at = Column(DateTime, default=datetime.datetime.utcnow, onupdate=datetime.datetime.utcnow)
    updated_by = Column(Integer, ForeignKey("users.id"), nullable=True)
    updater = relationship("User", backref="updated_settings")

# Drop all tables and recreate them
def recreate_database():
    Base.metadata.drop_all(bind=engine)
    Base.metadata.create_all(bind=engine)

# Check if database schema is up to date and migrate if needed
def ensure_database_schema():
    from sqlalchemy import text, inspect
    
    try:
        inspector = inspect(engine)
        
        # Check if users table has web_search_enabled column
        users_columns = [col['name'] for col in inspector.get_columns('users')]
        if 'web_search_enabled' not in users_columns:
            print("Adding web_search_enabled column to users table...")
            with engine.connect() as conn:
                conn.execute(text("ALTER TABLE users ADD COLUMN web_search_enabled INTEGER DEFAULT 1"))
                conn.commit()
            print("✅ web_search_enabled column added successfully")
        
        # Check if uploaded_csvs table has table_name column
        with engine.connect() as conn:
            result = conn.execute(text("SELECT table_name FROM uploaded_csvs LIMIT 1"))
    except Exception as e:
        print(f"Database schema outdated, recreating tables: {e}")
        recreate_database()

# Create tables and ensure schema is up to date
Base.metadata.create_all(bind=engine)
ensure_database_schema()

# Use explicit hash handlers to avoid dynamic handler import issues in frozen apps

SECRET_KEY = "supersecretkey"  # Change this in production
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60 * 24 * 7  # 1 week

def safe_string_conversion(value, default=""):
    """Safely convert any value to string, handling lists and other types"""
    if isinstance(value, str):
        return value
    elif isinstance(value, list):
        return ' '.join(str(item) for item in value if item)
    elif value is None:
        return default
    else:
        return str(value)

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

def is_model_enabled(model_name: str, db) -> bool:
    """Check if a model is enabled in the database"""
    try:
        model = db.query(AIModel).filter(AIModel.name == model_name, AIModel.is_enabled == 1).first()
        return model is not None
    except Exception:
        # Fallback to ALLOWED_OLLAMA_MODELS if database check fails
        return True

def get_server_setting(db, key: str, default: str = "") -> str:
    """Get a server setting value"""
    try:
        setting = db.query(ServerSettings).filter(ServerSettings.setting_key == key).first()
        return setting.setting_value if setting else default
    except Exception:
        return default

def set_server_setting(db, key: str, value: str, user_id: int = None):
    """Set or update a server setting"""
    try:
        setting = db.query(ServerSettings).filter(ServerSettings.setting_key == key).first()
        if setting:
            setting.setting_value = value
            setting.updated_at = datetime.datetime.utcnow()
            setting.updated_by = user_id
        else:
            setting = ServerSettings(setting_key=key, setting_value=value, updated_by=user_id)
            db.add(setting)
        db.commit()
        return True
    except Exception as e:
        print(f"Error setting server setting: {e}")
        db.rollback()
        return model_name in ALLOWED_OLLAMA_MODELS

def verify_password(plain_password, hashed_password):
    try:
        # If passlib handlers are available, prefer them
        if pbkdf2_sha256 is not None:
            try:
                if isinstance(hashed_password, str) and pbkdf2_sha256.identify(hashed_password):
                    return pbkdf2_sha256.verify(plain_password, hashed_password)
            except Exception:
                pass
        if bcrypt is not None:
            try:
                if isinstance(hashed_password, str) and bcrypt.identify(hashed_password):
                    return bcrypt.verify(plain_password, hashed_password)
            except Exception:
                pass

        # Fallback: support simple pbkdf2_sha256 string format created by this app
        if isinstance(hashed_password, str) and hashed_password.startswith("pbkdf2_sha256$"):
            parts = hashed_password.split("$")
            if len(parts) == 4:
                iterations = int(parts[1])
                salt = binascii.unhexlify(parts[2])
                dk = binascii.unhexlify(parts[3])
                newdk = hashlib.pbkdf2_hmac('sha256', plain_password.encode('utf-8'), salt, iterations)
                return hmac.compare_digest(newdk, dk)
        return False
    except Exception:
        return False

def get_password_hash(password):
    # Use passlib if available, otherwise produce a pbkdf2_sha256 formatted hash
    if pbkdf2_sha256 is not None:
        try:
            return pbkdf2_sha256.hash(password)
        except Exception:
            pass

    # Fallback implementation
    iterations = 120000
    salt = os.urandom(16)
    dk = hashlib.pbkdf2_hmac('sha256', password.encode('utf-8'), salt, iterations)
    return f"pbkdf2_sha256${iterations}${binascii.hexlify(salt).decode()}${binascii.hexlify(dk).decode()}"

def create_access_token(data: dict, expires_delta: datetime.timedelta | None = None):
    to_encode = data.copy()
    expire = datetime.datetime.utcnow() + (expires_delta or datetime.timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES))
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/login")

# Dependency to get current user
from pydantic import BaseModel as PydanticBaseModel
class TokenData(PydanticBaseModel):
    username: str | None = None
    role: str | None = None

def get_current_user(token: str = Depends(oauth2_scheme), db=Depends(get_db)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username = payload.get("sub")
        role = payload.get("role")
        if not isinstance(username, str) or username is None:
            raise credentials_exception
        token_data = TokenData(username=username, role=role)
    except JWTError:
        raise credentials_exception
    user = db.query(User).filter(User.username == token_data.username).first()
    if user is None:
        raise credentials_exception
    return user

def require_role(required_role: str):
    def role_checker(user: User = Depends(get_current_user)):
        if getattr(user, 'role', None) != required_role:
            raise HTTPException(status_code=403, detail="Insufficient permissions")
        return user
    return role_checker

# --- Registration Endpoint ---
class RegisterRequest(PydanticBaseModel):
    username: str
    password: str
    role: str = "user"  # Default role is 'user'

app = FastAPI()

# Add a set of allowed model names from ollama list
ALLOWED_OLLAMA_MODELS = {
    'anindya/prem1b-sql-ollama-fp116:latest',
    'llama3.1:latest',
    'qwen:0.5b',
    'gemma3:latest',
    'deepseek-v2:latest',
    'deepseek-coder:1.3b',
    'openchat:latest',
    'dolphin3:latest',
    'codellama:latest',
    'qwen2.5vl:latest',
    'deepseek-coder-v2:latest',
    'glm4:9b-chat-q4_0',
    'qwen3:0.6b',
    'llama3.2:1b',
    'deepseek-coder:latest',
    'llama3.1:8b',
    'nomic-embed-text:latest',
    'gpt-oss:latest',
    'gemma3:270m'
}

class PromptInput(BaseModel):
    prompt: str
    model: str
    session_id: str  # Added session_id for context
    use_web_search: bool = True  # Toggle for web search (default True)

class ChainRequest(BaseModel):
    prompt: str
    models: list[str]

# Data Analysis Models
class DataAnalysisRequest(BaseModel):
    prompt: str
    session_id: str
    model: str = "deepseek-coder-v2:latest"

class CSVUploadResponse(BaseModel):
    filename: str
    shape: tuple[int, int]  # Properly type the shape as a tuple of two integers
    columns: list[str]
    dtypes: dict[str, str]
    sample_data: list[dict]  # List of records from df.to_dict('records')

# Get API keys from environment variables
API_KEYS_STR = os.getenv("OPENROUTER_API_KEYS", "")
# Improved parsing to handle various formats
API_KEYS = []
if API_KEYS_STR:
    # Split by comma and handle various separators
    raw_keys = API_KEYS_STR.replace('\n', ',').replace(';', ',').split(',')
    for key in raw_keys:
        stripped_key = key.strip()
        if stripped_key and len(stripped_key) > 10:  # Basic validation for API key length
            API_KEYS.append(stripped_key)

print(f"🔑 Loaded {len(API_KEYS)} API keys from environment variable")
if len(API_KEYS) < 50:  # If you expect 50 keys but got fewer
    print(f"⚠️  Warning: Expected 50 keys but only loaded {len(API_KEYS)}")
    print(f"   Raw environment variable length: {len(API_KEYS_STR)}")
    print(f"   Raw preview: {API_KEYS_STR[:200]}...")

# Get allowed origins from environment variables
ALLOWED_ORIGINS = os.getenv("ALLOWED_ORIGINS", "http://localhost:3000,http://localhost:3001,http://127.0.0.1:3000,http://127.0.0.1:3001,http://localhost:8080,http://127.0.0.1:8080")
ALLOWED_ORIGINS_LIST = [origin.strip() for origin in ALLOWED_ORIGINS.split(",")]
# Always ensure dev origins are present
for dev_origin in ["http://localhost:8080", "http://127.0.0.1:8080"]:
    if dev_origin not in ALLOWED_ORIGINS_LIST:
        ALLOWED_ORIGINS_LIST.append(dev_origin)

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS_LIST,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

#BASE_URL = "https://oroute.ai/api/v1/chat/completions"

# ============================================================================
# DATA ANALYSIS FUNCTIONS (Converted from Streamlit)
# ============================================================================

# Ollama configuration for data analysis
OLLAMA_BASE_URL = "http://localhost:11434"
OLLAMA_URL = f"{OLLAMA_BASE_URL}/api/generate"
DATA_ANALYSIS_MODEL = "deepseek-coder-v2:latest"
MAX_RETRIES = 3
REQUEST_TIMEOUT = 300

def query_ollama_data_analysis(prompt, model=DATA_ANALYSIS_MODEL, timeout=REQUEST_TIMEOUT, retries=MAX_RETRIES):
    """Query Ollama API for data analysis with better error handling"""
    url = f"{OLLAMA_BASE_URL}/api/generate"
    headers = {"Content-Type": "application/json"}
    data = {
        "model": model,
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": 0.1,
            "top_p": 0.9,
            "num_ctx": 4096
        }
    }
    
    for attempt in range(retries):
        try:
            response = requests.post(
                url,
                headers=headers,
                json=data,
                timeout=timeout
            )
            
            if response.status_code == 200:
                result = response.json()
                if 'response' in result:
                    return result['response']
                else:
                    return "Error: Unexpected response format from Ollama"
            else:
                error_msg = f"HTTP {response.status_code}"
                try:
                    error_data = response.json()
                    if 'error' in error_data:
                        error_msg = error_data['error']
                except:
                    error_msg = response.text or error_msg
                
                if attempt == retries - 1:  # Last attempt
                    return f"API Error: {error_msg}"
                
        except requests.exceptions.Timeout:
            if attempt == retries - 1:  # Last attempt
                return "TIMEOUT_ERROR"
                
        except Exception as e:
            if attempt == retries - 1:  # Last attempt
                return f"Connection Error: {str(e)}"
        
        # Exponential backoff before retry
        if attempt < retries - 1:
            wait_time = (2 ** attempt) * 2  # 2, 4, 8 seconds
            time.sleep(wait_time)
    
    return "Failed after multiple attempts"

def generate_fallback_code(question, df):
    """
    Intelligent fallback code generator.
    Detects question intent (statistics, correlation, bar chart, histogram, scatter plot, etc.)
    and generates suitable Python code automatically.
    """
    import re
    question_lower = question.lower()
    numeric_cols = df.select_dtypes(include=['number']).columns.tolist()

    # --- 1️⃣ Detect if user asks for summary stats ---
    if any(k in question_lower for k in ["average", "mean", "minimum", "maximum", "max", "min", "median", "summary", "describe"]):
        for col in numeric_cols:
            if col.lower() in question_lower:
                return f"""
# Summary stats for {col}
avg = df["{col}"].mean()
min_val = df["{col}"].min()
max_val = df["{col}"].max()
median_val = df["{col}"].median()
print("Column: {col}")
print(f"Average (Mean): {{avg:.2f}}")
print(f"Median: {{median_val:.2f}}")
print(f"Minimum: {{min_val}}")
print(f"Maximum: {{max_val}}")
"""
        # If no specific column found
        return """
print("Summary Statistics for Numeric Columns:")
print(df.describe())
"""

    # --- 2️⃣ Detect histogram/distribution questions ---
    if any(k in question_lower for k in ["distribution", "histogram", "frequency", "spread"]):
        target_col = next((col for col in numeric_cols if col.lower() in question_lower), numeric_cols[0])
        return f"""
# Histogram for {target_col}
import matplotlib.pyplot as plt
plt.figure(figsize=(8,5))
plt.hist(df["{target_col}"].dropna(), bins=20, color='skyblue', edgecolor='black')
plt.title("Distribution of {target_col}")
plt.xlabel("{target_col}")
plt.ylabel("Frequency")
plt.tight_layout()
plt.show()
"""

    # --- 3️⃣ Detect correlation or "relationship" type questions ---
    if any(k in question_lower for k in ["correlation", "relationship", "compare", "association"]):
        if "heatmap" in question_lower or "matrix" in question_lower:
            return """
# Correlation heatmap
import matplotlib.pyplot as plt
import seaborn as sns
numeric_df = df.select_dtypes(include=['number'])
plt.figure(figsize=(10,8))
sns.heatmap(numeric_df.corr(), annot=True, cmap='coolwarm', fmt='.2f')
plt.title('Correlation Heatmap')
plt.tight_layout()
plt.show()
print(numeric_df.corr())
"""
        # Otherwise, scatter plot for "A vs B"
        matches = re.findall(r"([a-z_]+)\s*(?:vs|against|and)\s*([a-z_]+)", question_lower)
        if matches:
            col1, col2 = matches[0]
            col1 = next((c for c in df.columns if col1 in c.lower()), numeric_cols[0])
            col2 = next((c for c in df.columns if col2 in c.lower()), numeric_cols[1])
            return f"""
# Scatter plot between {col1} and {col2}
import matplotlib.pyplot as plt
plt.figure(figsize=(8,6))
plt.scatter(df["{col1}"], df["{col2}"], alpha=0.6, color='teal')
plt.title("{col1} vs {col2}")
plt.xlabel("{col1}")
plt.ylabel("{col2}")
plt.tight_layout()
plt.show()
"""

    # --- 4️⃣ Detect "by" or grouped average (bar chart) ---
    if "by" in question_lower or "group" in question_lower or "range" in question_lower:
        target_col = next((col for col in numeric_cols if col.lower() in question_lower and "by" not in col.lower()), "median_income")
        group_col = None
        for col in df.columns:
            if col.lower() in question_lower.split("by")[-1]:
                group_col = col
                break
        if not group_col:
            group_col = numeric_cols[1] if len(numeric_cols) > 1 else df.columns[0]

        return f"""
# Grouped bar chart of average {target_col} by {group_col} range
import matplotlib.pyplot as plt
import pandas as pd
df['{group_col}_range'] = pd.cut(df['{group_col}'], bins=5)
avg_data = df.groupby('{group_col}_range')['{target_col}'].mean()

plt.figure(figsize=(10,6))
avg_data.plot(kind='bar', color='cornflowerblue', edgecolor='black')
plt.title('Average {target_col} by {group_col} Range')
plt.xlabel('{group_col} Range')
plt.ylabel('Average {target_col}')
plt.xticks(rotation=45)
plt.tight_layout()
plt.show()
print(avg_data)
"""

    # --- 5️⃣ Detect trend / time-based questions ---
    if any(k in question_lower for k in ["trend", "over time", "year", "month", "timeline", "progression"]):
        target_col = next((col for col in numeric_cols if col.lower() in question_lower), numeric_cols[0])
        return f"""
# Line chart showing trend of {target_col} over index
import matplotlib.pyplot as plt
plt.figure(figsize=(8,5))
plt.plot(df.index, df["{target_col}"], color='orange')
plt.title("Trend of {target_col} over dataset index")
plt.xlabel("Index")
plt.ylabel("{target_col}")
plt.tight_layout()
plt.show()
"""

    # --- Default generic analysis ---
    return """
print("Dataset Shape:", df.shape)
print("\\nColumn Info:")
print(df.dtypes)
print("\\nBasic Statistics:")
print(df.describe())
print("\\nMissing Values:")
print(df.isnull().sum())
"""


def execute_data_analysis_code(code, df, table_name=None, globals_dict=None):
    """Safely execute Python code with dataframe context and SQL support"""
    # Ensure heavy data science libraries are loaded
    ensure_heavy_imports()
    
    if globals_dict is None:
        globals_dict = {}
    
    # Create SQLite engine if table_name is provided
    engine = None
    if table_name:
        engine = create_engine(f'sqlite:///databases/analysis.db')
    
    # Prepare execution environment with the actual dataframe
    exec_globals = {
        'df': df,
        'pd': pd,
        'np': np,
        'plt': plt,
        'sns': sns,
        'px': px,
        'go': go,
        'stats': stats,
        'LinearRegression': LinearRegression,
        'train_test_split': train_test_split,
        'r2_score': r2_score,
        'mean_squared_error': mean_squared_error,
        'execute_sql': lambda query: execute_sql_query(query, engine) if engine else None,
        'table_name': table_name,
        **globals_dict
    }
    
    # Clean the code - remove any file loading attempts
    cleaned_code = code.replace("pd.read_csv('path_to_your_dataset.csv')", "df")
    cleaned_code = cleaned_code.replace("pd.read_csv(", "# pd.read_csv(")
    cleaned_code = cleaned_code.replace("df = pd.read_csv", "# df = pd.read_csv")
    
    # Capture output
    output = io.StringIO()
    chart_data = None
    
    try:
        with redirect_stdout(output):
            exec(cleaned_code, exec_globals)
        
        # Get printed output
        printed_output = output.getvalue()
        
        # Check if a plot was created and save it as base64
        if plt.get_fignums():
            # Get the current figure
            fig = plt.gcf()
            
            # Save plot to base64 string
            buffer = io.BytesIO()
            fig.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
            buffer.seek(0)
            plot_data = buffer.getvalue()
            buffer.close()
            
            # Convert to base64
            chart_data = base64.b64encode(plot_data).decode('utf-8')
            
            # Clear the figure
            plt.clf()
            plt.close('all')
        
        return {
            'success': True,
            'output': printed_output,
            'chart': chart_data,
            'globals': exec_globals
        }
    
    except Exception as e:
        plt.close('all')  # Clean up any plots
        return {
            'success': False,
            'error': str(e),
            'traceback': traceback.format_exc()
        }

def save_to_sql_database(df, table_name, db_path="databases/analysis.db"):
    """Save DataFrame to SQLite database"""
    os.makedirs(os.path.dirname(db_path), exist_ok=True)
    
    # Create SQLite connection
    engine = create_engine(f'sqlite:///{db_path}')
    
    # Save DataFrame to SQL
    df.to_sql(table_name, engine, if_exists='replace', index=False)
    return engine

def execute_sql_query(query, engine):
    """Execute SQL query, print results, and return DataFrame"""
    try:
        df = pd.read_sql_query(query, engine)
        # Print a readable summary of the results
        print("SQL Query Results:")
        print(df.head(20).to_string(index=False))
        return df
    except Exception as e:
        print(f"SQL Query Error: {str(e)}")
        raise Exception(f"SQL Query Error: {str(e)}")
    
def get_database_engine():
    db_path = "databases/analysis.db"
    return create_engine(f"sqlite:///{db_path}")

@app.get("/list-tables")
def list_uploaded_tables(current_user: User = Depends(get_current_user)):
    engine = get_database_engine()
    inspector = inspect(engine)
    tables = inspector.get_table_names()
    if not tables:
        raise HTTPException(status_code=404, detail="No tables found in the database.")
    return {"tables": tables}

@app.get("/describe-table")
def describe_table(
    table_name: str = Query(..., description="Name of the table to describe"),
    current_user: User = Depends(get_current_user)
):
    engine = get_database_engine()
    inspector = inspect(engine)
    columns = inspector.get_columns(table_name)
    if not columns:
        raise HTTPException(status_code=404, detail=f"Table '{table_name}' not found.")
    return {
        "table": table_name,
        "columns": [
            {"name": col["name"], "type": str(col["type"]), "nullable": col["nullable"]}
            for col in columns
        ]
    }
@app.post("/execute-sql")
def execute_sql(
    query: str = Query(..., description="SQL query to execute (SELECT only)"),
    current_user: User = Depends(get_current_user)
):
    if not query.strip().lower().startswith("select"):
        raise HTTPException(status_code=400, detail="Only SELECT queries are allowed for safety.")

    engine = get_database_engine()
    try:
        df = pd.read_sql_query(query, engine)
        records = df.head(50).to_dict(orient="records")  # Limit rows for performance
        return {"success": True, "row_count": len(df), "columns": list(df.columns), "data": records}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"SQL execution error: {str(e)}")
    

def load_csv_file(file_path):
    """Load CSV file with multiple fallback methods"""
    import chardet
    
    # Read the file content first to detect encoding
    with open(file_path, 'rb') as f:
        raw_content = f.read()
    
    # Try to detect encoding
    try:
        detected_enc = chardet.detect(raw_content)['encoding']
        print(f"Detected encoding: {detected_enc}")
        if detected_enc is None:
            detected_enc = 'utf-8'
    except Exception as e:
        print(f"Error detecting encoding: {str(e)}")
        detected_enc = 'utf-8'
    
    # Try different methods to read the file
    methods = [
        (lambda: pd.read_csv(file_path, encoding=detected_enc), f"detected encoding ({detected_enc})"),
        (lambda: pd.read_csv(file_path, encoding='utf-8'), "utf-8"),
        (lambda: pd.read_csv(file_path, encoding='latin1'), "latin1"),
        (lambda: pd.read_csv(file_path, encoding='cp1252'), "cp1252"),
        (lambda: pd.read_csv(file_path, encoding='iso-8859-1'), "iso-8859-1"),
        (lambda: pd.read_csv(file_path, sep=';', encoding=detected_enc), f"semicolon separator, {detected_enc}"),
        (lambda: pd.read_csv(file_path, sep='\t', encoding=detected_enc), f"tab separator, {detected_enc}"),
        # Additional fallback methods for problematic files
        (lambda: pd.read_csv(file_path, encoding=detected_enc, on_bad_lines='skip'), f"skip bad lines, {detected_enc}"),
        (lambda: pd.read_csv(file_path, encoding=detected_enc, engine='python'), f"python engine, {detected_enc}"),
        (lambda: pd.read_csv(file_path, encoding=detected_enc, delimiter=None), f"auto delimiter, {detected_enc}")
    ]
    
    errors = []
    for method, description in methods:
        try:
            print(f"Trying to read CSV with {description}")
            df = method()
            if df is not None and not df.empty:
                print(f"Successfully read CSV with {description}")
                print(f"DataFrame shape: {df.shape}")
                print(f"Columns: {df.columns.tolist()}")
                
                # Clean column names
                df.columns = df.columns.str.strip()
                for col in df.columns:
                    try:
                        df[col] = df[col].replace(',', '', regex=True).astype(float)
                    except Exception:
                        pass  # ignore non-numeric columns

                print(f"Converted dtypes:\n{df.dtypes}")
                return df

        except Exception as e:
            error_msg = f"Method '{description}' failed: {str(e)}"
            print(error_msg)
            errors.append(error_msg)
            continue
    
    # If we get here, all methods failed
    error_details = "\n".join(errors)
    print(f"All methods failed to read the CSV file. Errors:\n{error_details}")
    return None

# ============================================================================
# EXISTING ENDPOINTS
# ============================================================================

# ✅ Health check endpoint
@app.get("/health")
def health_check():
    return {"status": "healthy", "message": "Backend is running"}

# ✅ API Keys diagnostic endpoint
@app.get("/debug/keys")
def debug_api_keys():
    """Diagnostic endpoint to check API key parsing"""
    raw_keys_str = os.getenv("OPENROUTER_API_KEYS", "")
    parsed_keys = [key.strip() for key in raw_keys_str.split(",") if key.strip()]
    
    return {
        "raw_environment_variable_length": len(raw_keys_str),
        "raw_environment_variable_preview": raw_keys_str[:100] + "..." if len(raw_keys_str) > 100 else raw_keys_str,
        "total_keys_after_parsing": len(parsed_keys),
        "parsed_keys_preview": [key[:10] + "..." + key[-4:] if len(key) > 14 else key for key in parsed_keys[:5]],
        "all_parsed_keys_count": len(parsed_keys),
        "environment_variable_name": "OPENROUTER_API_KEYS"
    }

# ✅ Check if web search is enabled globally and for the current user
@app.get("/settings/web-search-enabled")
def check_web_search_enabled(db=Depends(get_db), current_user: User = Depends(get_current_user)):
    """Check if web search is enabled globally and for the current user"""
    global_enabled = get_server_setting(db, "web_search_enabled", "true").lower() == "true"
    user_enabled = current_user.web_search_enabled == 1
    return {
        "enabled": global_enabled and user_enabled,
        "global_enabled": global_enabled,
        "user_enabled": user_enabled
    }

# ✅ Main endpoint to send prompt and receive model response
@app.post("/ask")
def ask_model(data: PromptInput, db=Depends(get_db), current_user: User = Depends(get_current_user, use_cache=False), skip_auth: bool = Query(False)):
    if skip_auth:
        # Guest mode for bots - skip user context/DB writes
        current_user = type('GuestUser', (), {'id': 'bot', 'web_search_enabled': 1})()

    # Fetch last 5 messages for context (oldest first) to include model switches
    history = db.query(ChatHistory)\
        .filter(ChatHistory.user_id == current_user.id, ChatHistory.session_id == data.session_id)\
        .order_by(ChatHistory.timestamp.desc())\
        .limit(5).all()

    context = ""
    previous_model = None
    for msg in reversed(history):
        # Check if there was a model switch and add transition context
        if msg.model_used and msg.model_used != data.model and msg.model_used != previous_model:
            context += f"[Note: Previous responses were from {msg.model_used} model]\n"
        context += f"User: {msg.prompt}\nAI: {msg.response}\n"
        previous_model = msg.model_used
    
    # Add current model context if switching from a different model
    if previous_model and previous_model != data.model:
        context += f"[Context: User has now switched to {data.model} model. Please maintain conversation continuity while applying your model's capabilities.]\n"

    # Fetch all extracted PDF texts for this session and user (in upload order)
    docs = db.query(UploadedDocument).filter(UploadedDocument.user_id == current_user.id, UploadedDocument.session_id == data.session_id, UploadedDocument.extracted_text != None).order_by(UploadedDocument.timestamp.asc()).all()

    pdf_context = ""
    for idx, doc in enumerate(docs):
        if doc.extracted_text:
            pdf_context += f"[PDF {idx+1} Content Start: {doc.filename}]\n{doc.extracted_text}\n[PDF {idx+1} Content End]\n"

    # Check if web search is enabled (global setting, user permission, and user toggle)
    web_search_enabled_global = get_server_setting(db, "web_search_enabled", "true").lower() == "true"
    user_has_permission = current_user.web_search_enabled == 1
    web_search_enabled = web_search_enabled_global and user_has_permission and data.use_web_search
    
    # Web Search with multiple fallbacks (only if enabled)
    searxng_context = ""
    if web_search_enabled:
        try:
            search_result = perform_web_search(data.prompt, max_results=5)
            if search_result["success"] and search_result["results"]:
                searxng_context = f"\n[WEB SEARCH RESULTS (via {search_result['source']}) START]\n"
                for idx, result in enumerate(search_result["results"], 1):
                    title = result.get("title", "No title")
                    content = result.get("content", "No content")
                    url = result.get("url", "No URL")
                    searxng_context += f"Source {idx}: {title}\nURL: {url}\nContent: {content}\n\n"
                searxng_context += "[WEB SEARCH RESULTS END]\n\n"
            else:
                print(f"Web search failed: {search_result.get('error', 'Unknown error')}")
        except Exception as e:
            print(f"Web search error: {e}")

    # Combine all contexts
    full_prompt = ""
    if pdf_context:
        full_prompt += pdf_context
    if searxng_context:
        full_prompt += searxng_context
    
    full_prompt += context + f"Based on the above search results and context, please answer the following question: {data.prompt}\nAI:"

    # Ensure Ollama is running (lazy initialization)
    ensure_ollama_running()
    
    # Use Ollama local API
    ollama_url = "http://localhost:11434/api/generate"
    payload = {
        "model": data.model,
        "prompt": full_prompt,
        "stream": True,
        "num_predict": 256,  # Increased for more detailed responses
        "temperature": 0.7,
    }

    # Check if model is enabled in database
    if not is_model_enabled(data.model, db):
        raise HTTPException(status_code=400, detail=f"Model '{data.model}' is not available or not enabled on this server.")

    def stream_ollama():
        # First, yield the search results as a separate message
        if searxng_context:
            yield json.dumps({
                "type": "search_results", 
                "content": searxng_context,
                "done": False
            }) + "\n"
        
        try:
            with requests.post(ollama_url, json=payload, stream=True, timeout=600) as response:
                response.raise_for_status()
                for line in response.iter_lines():
                    if line:
                        try:
                            chunk = line.decode()
                            data_json = json.loads(chunk)
                            resp = data_json.get("response", "")
                            done = data_json.get("done", False)
                            yield json.dumps({
                                "type": "model_response",
                                "response": resp, 
                                "done": done
                            }) + "\n"
                            if done:
                                break
                        except Exception:
                            continue
            # Ensure the last chunk is sent
            yield json.dumps({"type": "model_response", "done": True}) + "\n"
        except Exception as e:
            yield json.dumps({
                "type": "error",
                "error": f"Ollama error: {str(e)}", 
                "done": True
            }) + "\n"

    return StreamingResponse(stream_ollama(), media_type="application/jsonl")

# ============================================================================
# NEW DATA ANALYSIS ENDPOINTS
# ============================================================================

# Ensure upload directories exist - use backend directory for CSV uploads
CSV_UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "csv_uploads")
os.makedirs(CSV_UPLOAD_DIR, exist_ok=True)

from fastapi.middleware.cors import CORSMiddleware

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, replace with your frontend URL
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.post("/upload-csv")
async def upload_csv(
    file: UploadFile = File(...),
    session_id: str = Query(..., min_length=1, description="Session ID for tracking uploads"),
    db=Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Upload CSV file for data analysis"""
    # Ensure heavy data science libraries are loaded
    ensure_heavy_imports()
    
    print(f"Received upload request for session: {session_id}")
    
    # Validate file
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
        
    if not file.filename.lower().endswith('.csv'):
        raise HTTPException(
            status_code=400, 
            detail=f"Invalid file type: {file.filename}. Only CSV files are allowed"
        )
    
    # Create unique filename to avoid conflicts
    timestamp = int(time.time())
    filename = f"{current_user.id}_{session_id}_{timestamp}_{file.filename}"
    file_path = os.path.join(CSV_UPLOAD_DIR, filename)
    
    print(f"Processing file: {filename}")
    print(f"File size: {file.size if hasattr(file, 'size') else 'unknown'}")
    print(f"Content type: {file.content_type}")
    
    # Save file
    try:
        print("Reading file content...")
        content = await file.read()
        if not content:
            error_msg = "Uploaded file is empty"
            print(f"Error: {error_msg}")
            raise HTTPException(status_code=400, detail=error_msg)
        
        print(f"Read {len(content)} bytes, writing to {file_path}")
        
        # Ensure the directory exists
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        
        with open(file_path, "wb") as buffer:
            buffer.write(content)
        
        # Verify file was written
        if not os.path.exists(file_path):
            error_msg = f"File was not saved to {file_path}"
            print(f"Error: {error_msg}")
            raise HTTPException(status_code=500, detail=error_msg)
            
        file_size = os.path.getsize(file_path)
        print(f"File saved successfully. Size: {file_size} bytes")
        
        if file_size == 0:
            error_msg = "Saved file is empty"
            print(f"Error: {error_msg}")
            os.remove(file_path)  # Clean up empty file
            raise HTTPException(status_code=400, detail=error_msg)
            
    except Exception as e:
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(status_code=500, detail=f"Error saving file: {str(e)}")
    
    # Try to load and analyze the CSV
    try:
        print(f"Attempting to load CSV file: {file_path}")
        df = load_csv_file(file_path)
        if df is None or df.empty:
            if os.path.exists(file_path):
                os.remove(file_path)
            raise HTTPException(
                status_code=400, 
                detail="Could not read CSV file. The file might have an unsupported format or encoding. Please ensure it's a valid CSV file with UTF-8 or similar encoding."
            )
        
        # Get basic info about the dataset
        try:
            sample_data = df.head(5).to_dict('records') if len(df) > 0 else {}
            
            columns_info = {
                'columns': df.columns.tolist(),
                'dtypes': df.dtypes.astype(str).to_dict(),
                'shape': df.shape,
                'sample_data': sample_data
            }
        except Exception as e:
            if os.path.exists(file_path):
                os.remove(file_path)
            raise HTTPException(status_code=400, detail=f"Error processing CSV data: {str(e)}")
        
        # Store the CSV data in SQLite
        table_name = f"data_{current_user.id}_{session_id}_{int(time.time())}"
        engine = save_to_sql_database(df, table_name)
        
        # Store metadata in database
        csv_record = UploadedCSV(
            user_id=current_user.id,
            session_id=session_id,
            filename=file.filename,
            file_path=file_path,
            columns_info=json.dumps(columns_info),
            table_name=table_name
        )
        db.add(csv_record)
        db.commit()
        db.refresh(csv_record)
        
        # Create response using the Pydantic model
        response_data = CSVUploadResponse(
            filename=file.filename,
            shape=df.shape,
            columns=df.columns.tolist(),
            dtypes=df.dtypes.astype(str).to_dict(),
            sample_data=sample_data
        )
        
        return response_data
        
    except Exception as e:
        # Clean up file if processing failed
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(status_code=400, detail=f"Error processing CSV: {str(e)}")

@app.post("/data-analysis")
def data_analysis(
    data: DataAnalysisRequest,
    db=Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Process data analysis questions"""
    
    # Get the uploaded CSV for this session
    csv_record = db.query(UploadedCSV).filter(
        UploadedCSV.user_id == current_user.id,
        UploadedCSV.session_id == data.session_id
    ).order_by(UploadedCSV.timestamp.desc()).first()
    
    if not csv_record:
        raise HTTPException(status_code=400, detail="No CSV file uploaded for this session")
    
    if not os.path.exists(csv_record.file_path):
        raise HTTPException(status_code=400, detail="CSV file not found. Please re-upload.")
    
    # Load the CSV
    try:
        df = load_csv_file(csv_record.file_path)
        if df is None:
            raise HTTPException(status_code=400, detail="Could not load CSV file")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error loading CSV: {str(e)}")
    
    def stream_data_analysis():
        try:
            # Check if Ollama is available
            ollama_available = True
            try:
                response = requests.get(f"{OLLAMA_BASE_URL}/api/version", timeout=5)
                if response.status_code != 200:
                    ollama_available = False
            except:
                ollama_available = False
            
            if ollama_available and is_model_enabled(data.model, db):
                # Generate code using AI
                yield json.dumps({
                    "type": "status",
                    "message": "Generating analysis code with AI...",
                    "done": False
                }) + "\n"
                
                context = f"""
Dataset: {df.shape[0]} rows, {df.shape[1]} columns
Columns: {list(df.columns)}
Types: {dict(df.dtypes)}

RULES:
1. Use 'df' (already loaded)
2. Generate executable Python code only
3. Use print() for results
4. Keep code concise
5. Include matplotlib plots when appropriate

Question: {data.prompt}
Code:"""
                
                ai_response = query_ollama_data_analysis(context, model=data.model)
                
                if "Error" in ai_response or "TIMEOUT" in ai_response:
                    yield json.dumps({
                        "type": "status",
                        "message": "AI unavailable, using fallback code generation...",
                        "done": False
                    }) + "\n"
                    code = generate_fallback_code(data.prompt, df)
                else:
                    # Extract code from AI response
                    if '```python' in ai_response:
                        code_start = ai_response.find('```python') + 9
                        code_end = ai_response.find('```', code_start)
                        code = ai_response[code_start:code_end].strip()
                    else:
                        code = ai_response.strip()
            else:
                yield json.dumps({
                    "type": "status",
                    "message": "Using fallback code generation...",
                    "done": False
                }) + "\n"
                code = generate_fallback_code(data.prompt, df)
            
            # Send the generated code
            yield json.dumps({
                "type": "code",
                "content": code,
                "done": False
            }) + "\n"
            
            # Execute the code
            yield json.dumps({
                "type": "status",
                "message": "Executing analysis...",
                "done": False
            }) + "\n"
            
            result = execute_data_analysis_code(code, df, csv_record.table_name)
            
            if result['success']:
                # Send output if any
                if result.get('output'):
                    yield json.dumps({
                        "type": "output",
                        "content": result['output'],
                        "done": False
                    }) + "\n"
                
                # Send chart if any
                if result.get('chart'):
                    yield json.dumps({
                        "type": "chart",
                        "content": result['chart'],
                        "done": False
                    }) + "\n"
                
                # Send explanation
                explanation = f"Analysis completed successfully. The results show patterns and insights from your dataset with {df.shape} rows and {df.shape} columns."
                yield json.dumps({
                    "type": "explanation",
                    "content": explanation,
                    "done": False
                }) + "\n"
                
            else:
                yield json.dumps({
                    "type": "error",
                    "content": f"Execution error: {result['error']}",
                    "done": False
                }) + "\n"
            
            # Final done signal
            yield json.dumps({
                "type": "analysis_complete",
                "done": True
            }) + "\n"
            
        except Exception as e:
            yield json.dumps({
                "type": "error",
                "content": f"Analysis error: {str(e)}",
                "done": True
            }) + "\n"
    
    return StreamingResponse(stream_data_analysis(), media_type="application/jsonl")

@app.get("/csv-info")
def get_csv_info(
    session_id: str = Query(...),
    db=Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Get information about uploaded CSV for this session"""
    csv_record = db.query(UploadedCSV).filter(
        UploadedCSV.user_id == current_user.id,
        UploadedCSV.session_id == session_id
    ).order_by(UploadedCSV.timestamp.desc()).first()
    
    if not csv_record:
        return {"has_csv": False}
    
    columns_info = json.loads(csv_record.columns_info) if csv_record.columns_info else {}
    
    return {
        "has_csv": True,
        "filename": csv_record.filename,
        "uploaded_at": csv_record.timestamp.isoformat(),
        **columns_info
    }

# ============================================================================
# EXISTING ENDPOINTS (CONTINUED)
# ============================================================================

@app.get("/search/searxng")
def searxng_search(query: str = Query(..., description="Search query"), max_results: int = 5):
    """Search using multiple search engines with fallbacks (DuckDuckGo, SearXNG instances)."""
    try:
        search_result = perform_web_search(query, max_results=max_results)
        
        if search_result["success"]:
            return {
                "results": search_result["results"],
                "query": query,
                "source": search_result["source"],
                "success": True
            }
        else:
            return {
                "error": search_result.get("error", "All search methods failed"),
                "results": [],
                "query": query,
                "success": False
            }
    except Exception as e:
        return {
            "error": str(e),
            "results": [],
            "query": query,
            "success": False
        }

@app.post("/chain")
def chain_models(data: ChainRequest, db=Depends(get_db)):
    # Ensure Ollama is running (lazy initialization)
    ensure_ollama_running()
    
    def stream_chain():
        current_prompt = data.prompt
        for model_id in data.models:
            if not is_model_enabled(model_id, db):
                yield json.dumps({"model": model_id, "response": f"Model '{model_id}' is not available or not enabled on this server.", "done": True}) + "\n"
                current_prompt = data.prompt
                continue
            ollama_url = "http://localhost:11434/api/generate"
            payload = {
                "model": model_id,
                "prompt": current_prompt,
                "stream": True,
                "num_predict": 64,
                "temperature": 0.8,
            }
            full_response = ""
            try:
                with requests.post(ollama_url, json=payload, stream=True, timeout=600) as response:
                    response.raise_for_status()
                    for line in response.iter_lines():
                        if line:
                            try:
                                chunk = line.decode()
                                data_json = json.loads(chunk)
                                resp = data_json.get("response", "")
                                done = data_json.get("done", False)
                                full_response += resp
                                yield json.dumps({"model": model_id, "response": resp, "done": False}) + "\n"
                                if done:
                                    break
                            except Exception:
                                continue
                    # After model is done, send done for this model
                    yield json.dumps({"model": model_id, "done": True}) + "\n"
                    # Pass the full output to the next model
                    current_prompt = full_response
            except Exception as e:
                yield json.dumps({"model": model_id, "response": f"Ollama error: {str(e)}", "done": True}) + "\n"
                current_prompt = data.prompt
    return StreamingResponse(stream_chain(), media_type="application/jsonl")

@app.post("/upload")
def upload_files(files: List[UploadFile] = File(...), session_id: str = Query(...), db=Depends(get_db), current_user: User = Depends(get_current_user)):
    # Ensure file processing libraries are loaded
    ensure_file_processing_imports()
    
    saved_files = []
    for file in files:
        if file.filename is None:
            continue
        filename = file.filename
        file_path = os.path.join(UPLOAD_DIR, filename)
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        extracted_text = None
        
        # Handle file types
        if filename.lower().endswith(".pdf"):
            try:
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() or ""
                    extracted_text = text
            except Exception as e:
                extracted_text = None
        elif filename.lower().endswith(".docx"):
            try:
                docx_file = docx.Document(file_path)
                text = "\n".join([para.text for para in docx_file.paragraphs])
                extracted_text = text
            except Exception as e:
                extracted_text = None
        elif filename.lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".gif")):
            try:
                image = Image.open(file_path)
                text = pytesseract.image_to_string(image)
                extracted_text = text
            except Exception as e:
                extracted_text = None
        
        # Store in UploadedDocument
        doc = UploadedDocument(
            user_id=current_user.id,
            session_id=session_id,
            filename=filename,
            extracted_text=extracted_text
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)
        saved_files.append({"filename": filename, "extracted_text": bool(extracted_text)})
    return {"files": saved_files}

# ============================================================================
# OCR ENDPOINTS
# ============================================================================

from pydantic import BaseModel

class ChatRequest(BaseModel):
    user_id: str
    message: str

@app.post("/chat")
async def discord_chat(request: ChatRequest):
    """Discord bot chat endpoint - simple non-streaming Ollama proxy."""
    session_id = request.user_id
    prompt = request.message
    model = "gemma3:270m"  # Default model
    
    ensure_ollama_running()
    ollama_url = "http://localhost:11434/api/generate"
    payload = {
        "model": model,
        "prompt": f"User: {prompt}\\nAI:",
        "stream": False,
        "options": {"temperature": 0.7, "num_predict": 512}
    }
    try:
        resp = requests.post(ollama_url, json=payload, timeout=30)
        print(f"DEBUG Ollama payload: model={model}")  # Debug
        resp.raise_for_status()
        data = resp.json()
        reply = data.get("response", "No response generated.").strip()
        return {"response": reply[:1997] + "..." if len(reply) > 1997 else reply}
    except Exception as e:
        return {"response": f"❌ AI Error: {str(e)}"}


@app.get("/ocr/languages")
def get_ocr_languages():
    """Get list of supported OCR languages"""
    return {
        "languages": SUPPORTED_OCR_LANGUAGES,
        "total": len(SUPPORTED_OCR_LANGUAGES)
    }

@app.post("/ocr")
async def perform_ocr_endpoint(
    file: UploadFile = File(...),
    languages: str = Query(default="hin", description="Language code(s) for OCR (e.g., 'eng', 'hin', 'hin+eng')"),
    enhance: bool = Query(default=True, description="Enable image preprocessing"),
    output_format: str = Query(default="text", description="Output format: text, json, or both"),
    db=Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Perform OCR on an image file with multilingual support"""
    # Ensure file processing libraries are loaded
    ensure_file_processing_imports()
    
    # Validate file type
    if not file.filename.lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".gif", ".webp")):
        raise HTTPException(
            status_code=400, 
            detail="Invalid file type. Only image files are supported for OCR."
        )
    
    # Validate language codes
    requested_langs = languages.split("+")
    unsupported = [lang for lang in requested_langs if lang not in SUPPORTED_OCR_LANGUAGES]
    if unsupported:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported language(s): {', '.join(unsupported)}. Supported: {', '.join(SUPPORTED_OCR_LANGUAGES.keys())}"
        )
    
    # Save uploaded file temporarily
    import tempfile
    with tempfile.NamedTemporaryFile(delete=False, suffix=Path(file.filename).suffix) as tmp:
        content = await file.read()
        tmp.write(content)
        tmp_path = tmp.name
    
    try:
        # Perform OCR
        result = perform_multilingual_ocr(tmp_path, languages=languages, enhance=enhance)
        
        if not result["success"]:
            raise HTTPException(status_code=500, detail=f"OCR failed: {result.get('error', 'Unknown error')}")
        
        # Prepare response based on format
        response = {}
        
        if output_format in ["text", "both"]:
            response["text"] = result["text"]
        
        if output_format in ["json", "both"]:
            response["ocr_data"] = {
                "text": result["text"],
                "confidence": result["confidence"],
                "languages_used": result["languages_used"],
                "success": result["success"]
            }
        
        response["filename"] = file.filename
        response["languages"] = languages
        response["enhance"] = enhance
        
        return response
        
    finally:
        # Clean up temp file
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

@app.post("/ocr/batch")
async def perform_batch_ocr(
    files: List[UploadFile] = File(...),
    languages: str = Query(default="hin", description="Language code(s) for OCR"),
    enhance: bool = Query(default=True, description="Enable image preprocessing"),
    db=Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Perform OCR on multiple images"""
    # Ensure file processing libraries are loaded
    ensure_file_processing_imports()
    
    import tempfile
    
    results = []
    
    for file in files:
        # Validate file type
        if not file.filename.lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".gif", ".webp")):
            results.append({
                "filename": file.filename,
                "success": False,
                "error": "Invalid file type"
            })
            continue
        
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(file.filename).suffix) as tmp:
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name
        
        try:
            # Perform OCR
            result = perform_multilingual_ocr(tmp_path, languages=languages, enhance=enhance)
            
            results.append({
                "filename": file.filename,
                "success": result["success"],
                "text": result["text"] if result["success"] else None,
                "confidence": result["confidence"] if result["success"] else None,
                "error": result.get("error")
            })
        except Exception as e:
            results.append({
                "filename": file.filename,
                "success": False,
                "error": str(e)
            })
        finally:
            # Clean up temp file
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
    
    return {"results": results, "total": len(files), "successful": sum(1 for r in results if r["success"])}

# Continue with remaining endpoints...
EXA_API_KEY = os.getenv("EXA_API_KEY", "")

@app.get("/search/web")
def exa_web_search(query: str = Query(..., description="Search query")):
    if not EXA_API_KEY:
        return {"error": "EXA_API_KEY not set in environment variables."}
    url = "https://api.exa.ai/search"
    headers = {
        "Accept": "application/json",
        "Content-Type": "application/json",
        "Authorization": f"Bearer {EXA_API_KEY}"
    }
    payload = {
        "query": query,
        "numResults": 5
    }
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=15)
        response.raise_for_status()
        return response.json()
    except Exception as e:
        return {"error": str(e)}

TAVILY_API_KEY = os.getenv("TAVILY_API_KEY", "")

@app.get("/search/youtube")
def tavily_youtube_search(query: str = Query(..., description="Search query"), max_results: int = 5):
    """Search YouTube videos using Tavily API."""
    if not TAVILY_API_KEY:
        return {"error": "TAVILY_API_KEY not set in environment variables."}
    url = "https://api.tavily.com/search"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {TAVILY_API_KEY}"
    }
    payload = {
        "query": query,
        "search_depth": "basic",
        "include_answer": False,
        "include_images": False,
        "include_raw_content": False,
        "max_results": max_results,
        "include_sources": True,
        "search_type": "youtube"
    }
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=15)
        print("Tavily YouTube response:", response.text)  # Debug print
        response.raise_for_status()
        return response.json()
    except Exception as e:
        return {"error": str(e)}

@app.get("/search/reddit")
def tavily_reddit_search(query: str = Query(..., description="Search query"), max_results: int = 5):
    """Search Reddit posts using Tavily API."""
    if not TAVILY_API_KEY:
        return {"error": "TAVILY_API_KEY not set in environment variables."}
    url = "https://api.tavily.com/search"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {TAVILY_API_KEY}"
    }
    payload = {
        "query": query,
        "search_depth": "basic",
        "include_answer": False,
        "include_images": False,
        "include_raw_content": False,
        "max_results": max_results,
        "include_sources": True,
        "search_type": "reddit"
    }
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=15)
        print("Tavily Reddit response:", response.text)  # Debug print
        response.raise_for_status()
        return response.json()
    except Exception as e:
        return {"error": str(e)}

@app.get("/search/academic")
def tavily_academic_search(query: str = Query(..., description="Search query"), max_results: int = 5):
    """Search academic papers using Tavily API."""
    if not TAVILY_API_KEY:
        return {"error": "TAVILY_API_KEY not set in environment variables."}
    url = "https://api.tavily.com/search"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {TAVILY_API_KEY}"
    }
    payload = {
        "query": query,
        "search_depth": "basic",
        "include_answer": False,
        "include_images": False,
        "include_raw_content": False,
        "max_results": max_results,
        "include_sources": True,
        "search_type": "academic"
    }
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=15)
        print("Tavily Academic response:", response.text)  # Debug print
        response.raise_for_status()
        return response.json()
    except Exception as e:
        return {"error": str(e)}

@app.get("/search/crypto")
def tavily_crypto_search(query: str = Query(..., description="Search query"), max_results: int = 5):
    """Search crypto news using Tavily API."""
    if not TAVILY_API_KEY:
        return {"error": "TAVILY_API_KEY not set in environment variables."}
    url = "https://api.tavily.com/search"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {TAVILY_API_KEY}"
    }
    payload = {
        "query": query,
        "search_depth": "basic",
        "include_answer": False,
        "include_images": False,
        "include_raw_content": False,
        "max_results": max_results,
        "include_sources": True,
        "search_type": "crypto"
    }
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=15)
        print("Tavily Crypto response:", response.text)  # Debug print
        response.raise_for_status()
        return response.json()
    except Exception as e:
        return {"error": str(e)}

# --- Registration Endpoint ---
@app.post("/register")
def register_user(data: RegisterRequest, db=Depends(get_db)):
    if db.query(User).filter(User.username == data.username).first():
        raise HTTPException(status_code=400, detail="Username already registered")
    user = User(
        username=data.username,
        password_hash=get_password_hash(data.password),
        role=data.role,
        created_at=datetime.datetime.utcnow()  # Set creation time
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    return {"msg": "User registered successfully"}

# --- Login Endpoint ---
class Token(PydanticBaseModel):
    access_token: str
    token_type: str

@app.post("/login", response_model=Token)
def login(form_data: OAuth2PasswordRequestForm = Depends(), db=Depends(get_db)):
    user = db.query(User).filter(User.username == form_data.username).first()
    if not user or not verify_password(form_data.password, user.password_hash):
        raise HTTPException(status_code=400, detail="Incorrect username or password")
    access_token = create_access_token(data={"sub": user.username, "role": user.role})
    return {"access_token": access_token, "token_type": "bearer"}

# --- Get Current User ---
@app.get("/me")
def read_users_me(current_user: User = Depends(get_current_user)):
    return {"username": current_user.username, "role": current_user.role}

# Example of RBAC-protected endpoint
@app.get("/admin-only")
def admin_only_endpoint(user: User = Depends(require_role("admin"))):
    return {"msg": f"Hello, admin {user.username}!"}

# ============================================================================
# ADMIN AI MODEL MANAGEMENT ENDPOINTS
# ============================================================================

class AIModelCreate(BaseModel):
    name: str
    description: str | None = None
    usecases: str | None = None
    is_enabled: int = 1
    is_default: int = 0

class AIModelUpdate(BaseModel):
    name: str | None = None
    description: str | None = None
    usecases: str | None = None
    is_enabled: int | None = None
    is_default: int | None = None

class AIModelResponse(BaseModel):
    id: int
    name: str
    description: str | None
    usecases: str | None = None
    is_enabled: int
    is_default: int
    created_by: int | None
    created_at: datetime.datetime

# Get all AI models (admin only)
@app.get("/admin/models", response_model=list[AIModelResponse])
def get_all_models(
    db=Depends(get_db),
    current_user: User = Depends(require_role("admin"))
):
    models = db.query(AIModel).order_by(AIModel.name).all()
    return [
        AIModelResponse(
            id=model.id,
            name=model.name,
            description=model.description,
            usecases=model.usecases,
            is_enabled=model.is_enabled,
            is_default=model.is_default,
            created_by=model.created_by,
            created_at=model.created_at
        )
        for model in models
    ]

# Create new AI model (admin only)
@app.post("/admin/models", response_model=AIModelResponse)
def create_model(
    data: AIModelCreate,
    db=Depends(get_db),
    current_user: User = Depends(require_role("admin"))
):
    # Check if model with same name exists
    existing = db.query(AIModel).filter(AIModel.name == data.name).first()
    if existing:
        raise HTTPException(status_code=400, detail="Model with this name already exists")
    
    # If setting as default, unset other defaults
    if data.is_default == 1:
        db.query(AIModel).filter(AIModel.is_default == 1).update({"is_default": 0})
    
    model = AIModel(
        name=data.name,
        description=data.description,
        is_enabled=data.is_enabled,
        is_default=data.is_default,
        created_by=current_user.id
    )
    db.add(model)
    db.commit()
    db.refresh(model)
    return AIModelResponse(
        id=model.id,
        name=model.name,
        description=model.description,
        is_enabled=model.is_enabled,
        is_default=model.is_default,
        created_by=model.created_by,
        created_at=model.created_at
    )

# Get single model by ID (admin only)
@app.get("/admin/models/{model_id}", response_model=AIModelResponse)
def get_model(
    model_id: int,
    db=Depends(get_db),
    current_user: User = Depends(require_role("admin"))
):
    model = db.query(AIModel).filter(AIModel.id == model_id).first()
    if not model:
        raise HTTPException(status_code=404, detail="Model not found")
    return AIModelResponse(
        id=model.id,
        name=model.name,
        description=model.description,
        is_enabled=model.is_enabled,
        is_default=model.is_default,
        created_by=model.created_by,
        created_at=model.created_at
    )

# Update model (admin only)
@app.put("/admin/models/{model_id}", response_model=AIModelResponse)
def update_model(
    model_id: int,
    data: AIModelUpdate,
    db=Depends(get_db),
    current_user: User = Depends(require_role("admin"))
):
    model = db.query(AIModel).filter(AIModel.id == model_id).first()
    if not model:
        raise HTTPException(status_code=404, detail="Model not found")
    
    # If setting as default, unset other defaults
    if data.is_default == 1 and model.is_default == 0:
        db.query(AIModel).filter(AIModel.is_default == 1).update({"is_default": 0})
    
    # Update fields if provided
    if data.name is not None:
        # Check for duplicate name
        existing = db.query(AIModel).filter(AIModel.name == data.name, AIModel.id != model_id).first()
        if existing:
            raise HTTPException(status_code=400, detail="Model with this name already exists")
        model.name = data.name
    if data.description is not None:
        model.description = data.description
    if data.is_enabled is not None:
        model.is_enabled = data.is_enabled
    if data.is_default is not None:
        model.is_default = data.is_default
    
    db.commit()
    db.refresh(model)
    return AIModelResponse(
        id=model.id,
        name=model.name,
        description=model.description,
        is_enabled=model.is_enabled,
        is_default=model.is_default,
        created_by=model.created_by,
        created_at=model.created_at
    )

# Delete model (admin only)
@app.delete("/admin/models/{model_id}")
def delete_model(
    model_id: int,
    db=Depends(get_db),
    current_user: User = Depends(require_role("admin"))
):
    model = db.query(AIModel).filter(AIModel.id == model_id).first()
    if not model:
        raise HTTPException(status_code=404, detail="Model not found")
    
    db.delete(model)
    db.commit()
    return {"msg": "Model deleted successfully"}

# Toggle model enabled status (admin only)
@app.post("/admin/models/{model_id}/toggle")
def toggle_model(
    model_id: int,
    db=Depends(get_db),
    current_user: User = Depends(require_role("admin"))
):
    model = db.query(AIModel).filter(AIModel.id == model_id).first()
    if not model:
        raise HTTPException(status_code=404, detail="Model not found")
    
    model.is_enabled = 1 - model.is_enabled  # Toggle between 0 and 1
    db.commit()
    return {
        "msg": f"Model {'enabled' if model.is_enabled else 'disabled'} successfully",
        "is_enabled": model.is_enabled
    }

# Set model as default (admin only)
@app.post("/admin/models/{model_id}/set-default")
def set_default_model(
    model_id: int,
    db=Depends(get_db),
    current_user: User = Depends(require_role("admin"))
):
    model = db.query(AIModel).filter(AIModel.id == model_id).first()
    if not model:
        raise HTTPException(status_code=404, detail="Model not found")
    
    # Unset all other defaults
    db.query(AIModel).filter(AIModel.is_default == 1).update({"is_default": 0})
    
    # Set this model as default
    model.is_default = 1
    model.is_enabled = 1  # Ensure default model is enabled
    db.commit()
    
    return {"msg": f"Model '{model.name}' set as default"}

# Get only enabled models (for dropdown - all authenticated users)
@app.get("/models/enabled")
def get_enabled_models(
    db=Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    models = db.query(AIModel).filter(AIModel.is_enabled == 1).order_by(AIModel.name).all()
    return [
        {
            "id": model.id,
            "name": model.name,
            "description": model.description,
            "is_default": model.is_default
        }
        for model in models
    ]

# Sync models with ALLOWED_OLLAMA_MODELS (admin only)
@app.post("/admin/models/sync-ollama")
def sync_ollama_models(
    db=Depends(get_db),
    current_user: User = Depends(require_role("admin"))
):
    """Sync database models with ALLOWED_OLLAMA_MODELS set"""
    synced = []
    for model_name in ALLOWED_OLLAMA_MODELS:
        existing = db.query(AIModel).filter(AIModel.name == model_name).first()
        if not existing:
            model = AIModel(
                name=model_name,
                description=f"Auto-synced from Ollama: {model_name}",
                is_enabled=1,
                is_default=0,
                created_by=current_user.id
            )
            db.add(model)
            synced.append(model_name)
    
    db.commit()
    return {
        "msg": f"Synced {len(synced)} new models from ALLOWED_OLLAMA_MODELS",
        "synced_models": synced,
        "total_models": len(ALLOWED_OLLAMA_MODELS)
    }

# Pull model from Ollama (admin only)
class PullModelRequest(BaseModel):
    model_name: str

@app.post("/admin/models/pull-ollama")
def pull_ollama_model(
    data: PullModelRequest,
    db=Depends(get_db),
    current_user: User = Depends(require_role("admin"))
):
    """Pull a model from Ollama registry"""
    # Ensure Ollama is running (lazy initialization)
    ensure_ollama_running()
    
    try:
        model_name = data.model_name.strip()
        if not model_name:
            raise HTTPException(status_code=400, detail="Model name is required")
        
        # Call Ollama API to pull the model
        ollama_url = "http://localhost:11434/api/pull"
        
        def stream_pull():
            try:
                with requests.post(
                    ollama_url, 
                    json={"name": model_name}, 
                    stream=True, 
                    timeout=1800  # 30 minutes timeout for large models
                ) as response:
                    if response.status_code != 200:
                        yield json.dumps({
                            "error": f"Ollama API error: {response.status_code}",
                            "status": "error"
                        }) + "\n"
                        return
                    
                    for line in response.iter_lines():
                        if line:
                            try:
                                data = json.loads(line)
                                yield json.dumps(data) + "\n"
                                
                                # If pull is complete, add to database
                                if data.get("status") == "success" or "successfully" in data.get("status", "").lower():
                                    # Check if model already exists
                                    existing = db.query(AIModel).filter(AIModel.name == model_name).first()
                                    if not existing:
                                        model = AIModel(
                                            name=model_name,
                                            description=f"Pulled from Ollama: {model_name}",
                                            is_enabled=1,
                                            is_default=0,
                                            created_by=current_user.id
                                        )
                                        db.add(model)
                                        db.commit()
                            except json.JSONDecodeError:
                                continue
            except Exception as e:
                yield json.dumps({
                    "error": f"Pull error: {str(e)}",
                    "status": "error"
                }) + "\n"
        
        return StreamingResponse(stream_pull(), media_type="application/x-ndjson")
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to pull model: {str(e)}")

# ============================================================================
# ADMIN USER MANAGEMENT ENDPOINTS
# ============================================================================

class UserResponse(BaseModel):
    id: int
    username: str
    role: str
    web_search_enabled: int
    created_at: datetime.datetime

class UserRoleUpdate(BaseModel):
    role: str

# Get all users (admin only)
@app.get("/admin/users", response_model=list[UserResponse])
def get_all_users(
    db=Depends(get_db),
    current_user: User = Depends(require_role("admin"))
):
    users = db.query(User).order_by(User.created_at.desc()).all()
    return [
        UserResponse(
            id=user.id,
            username=user.username,
            role=user.role,
            web_search_enabled=user.web_search_enabled if hasattr(user, 'web_search_enabled') else 1,
            created_at=user.created_at
        )
        for user in users
    ]

# Update user role (admin only)
@app.put("/admin/users/{user_id}/role", response_model=UserResponse)
def update_user_role(
    user_id: int,
    data: UserRoleUpdate,
    db=Depends(get_db),
    current_user: User = Depends(require_role("admin"))
):
    # Don't allow admin to change their own role
    if user_id == current_user.id:
        raise HTTPException(status_code=400, detail="Cannot change your own role")
    
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Validate role
    if data.role not in ["user", "admin"]:
        raise HTTPException(status_code=400, detail="Invalid role. Must be 'user' or 'admin'")
    
    user.role = data.role
    db.commit()
    db.refresh(user)
    
    return UserResponse(
        id=user.id,
        username=user.username,
        role=user.role,
        web_search_enabled=user.web_search_enabled if hasattr(user, 'web_search_enabled') else 1,
        created_at=user.created_at
    )

# Toggle user web search permission (admin only)
@app.post("/admin/users/{user_id}/toggle-web-search")
def toggle_user_web_search(
    user_id: int,
    db=Depends(get_db),
    current_user: User = Depends(require_role("admin"))
):
    """Toggle web search permission for a specific user"""
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Toggle the web_search_enabled field
    user.web_search_enabled = 0 if user.web_search_enabled == 1 else 1
    db.commit()
    db.refresh(user)
    
    status = "enabled" if user.web_search_enabled == 1 else "disabled"
    return {
        "msg": f"Web search {status} for user '{user.username}'",
        "web_search_enabled": user.web_search_enabled
    }

# Delete user (admin only)
@app.delete("/admin/users/{user_id}")
def delete_user(
    user_id: int,
    db=Depends(get_db),
    current_user: User = Depends(require_role("admin"))
):
    # Don't allow admin to delete themselves
    if user_id == current_user.id:
        raise HTTPException(status_code=400, detail="Cannot delete your own account")
    
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    username = user.username
    db.delete(user)
    db.commit()
    
    return {"msg": f"User '{username}' deleted successfully"}

# Get user stats (admin only)
@app.get("/admin/stats")
def get_admin_stats(
    db=Depends(get_db),
    current_user: User = Depends(require_role("admin"))
):
    total_users = db.query(User).count()
    total_admins = db.query(User).filter(User.role == "admin").count()
    total_regular_users = db.query(User).filter(User.role == "user").count()
    total_chats = db.query(ChatHistory).count()
    total_documents = db.query(UploadedDocument).count()
    total_csvs = db.query(UploadedCSV).count()
    total_models = db.query(AIModel).count()
    enabled_models = db.query(AIModel).filter(AIModel.is_enabled == 1).count()
    
    return {
        "users": {
            "total": total_users,
            "admins": total_admins,
            "regular": total_regular_users
        },
        "content": {
            "chats": total_chats,
            "documents": total_documents,
            "csv_files": total_csvs
        },
        "models": {
            "total": total_models,
            "enabled": enabled_models
        }
    }

# Get server settings (admin only)
@app.get("/admin/settings")
def get_server_settings(
    db=Depends(get_db),
    current_user: User = Depends(require_role("admin"))
):
    """Get all server settings"""
    try:
        settings = db.query(ServerSettings).all()
        settings_dict = {s.setting_key: s.setting_value for s in settings}
        
        # Add default settings if they don't exist
        if "web_search_enabled" not in settings_dict:
            settings_dict["web_search_enabled"] = "true"
        
        return {
            "settings": settings_dict,
            "success": True
        }
    except Exception as e:
        return {"error": str(e), "success": False}

# Update server setting (admin only)
@app.post("/admin/settings/{setting_key}")
def update_server_setting(
    setting_key: str,
    value: dict,
    db=Depends(get_db),
    current_user: User = Depends(require_role("admin"))
):
    """Update a server setting"""
    try:
        setting_value = value.get("value", "")
        set_server_setting(db, setting_key, setting_value, current_user.id)
        return {
            "msg": f"Setting '{setting_key}' updated to '{setting_value}'",
            "success": True
        }
    except Exception as e:
        return {"error": str(e), "success": False}

# ============================================================================
# DATABASE CONNECTION MANAGEMENT ENDPOINTS
# ============================================================================

from db_connectors import DatabaseConnector
from routes.tools import register_tools_routes

register_tools_routes(app)

class DBConnectionCreate(BaseModel):
    name: str
    db_type: str
    host: str | None = None
    port: int | None = None
    database: str
    username: str | None = None
    password: str | None = None
    connection_string: str | None = None

class DBConnectionResponse(BaseModel):
    id: int
    name: str
    db_type: str
    host: str | None
    port: int | None
    database: str
    username: str | None
    is_active: int
    created_at: datetime.datetime
    last_used: datetime.datetime | None

class DBConnectionTest(BaseModel):
    db_type: str
    host: str | None = None
    port: int | None = None
    database: str
    username: str | None = None
    password: str | None = None
    connection_string: str | None = None

# Get all database connections for current user
@app.get("/db-connections", response_model=list[DBConnectionResponse])
def get_db_connections(
    db=Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    from models import DatabaseConnection
    connections = db.query(DatabaseConnection).filter(
        DatabaseConnection.user_id == current_user.id
    ).order_by(DatabaseConnection.created_at.desc()).all()
    
    return [
        DBConnectionResponse(
            id=conn.id,
            name=conn.name,
            db_type=conn.db_type,
            host=conn.host,
            port=conn.port,
            database=conn.database,
            username=conn.username,
            is_active=conn.is_active,
            created_at=conn.created_at,
            last_used=conn.last_used
        )
        for conn in connections
    ]

# Create new database connection
@app.post("/db-connections", response_model=DBConnectionResponse)
def create_db_connection(
    data: DBConnectionCreate,
    db=Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    from models import DatabaseConnection
    
    # Build connection string
    try:
        conn_string = DatabaseConnector.build_connection_string(
            db_type=data.db_type,
            host=data.host,
            port=data.port,
            database=data.database,
            username=data.username,
            password=data.password,
            custom_string=data.connection_string
        )
        
        # Test connection
        success, message = DatabaseConnector.test_connection(conn_string)
        if not success:
            raise HTTPException(status_code=400, detail=message)
        
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))
    
    # Save connection
    connection = DatabaseConnection(
        user_id=current_user.id,
        name=data.name,
        db_type=data.db_type,
        host=data.host,
        port=data.port,
        database=data.database,
        username=data.username,
        password=data.password,  # In production, encrypt this!
        connection_string=conn_string,
        is_active=1
    )
    
    db.add(connection)
    db.commit()
    db.refresh(connection)
    
    return DBConnectionResponse(
        id=connection.id,
        name=connection.name,
        db_type=connection.db_type,
        host=connection.host,
        port=connection.port,
        database=connection.database,
        username=connection.username,
        is_active=connection.is_active,
        created_at=connection.created_at,
        last_used=connection.last_used
    )

# Test database connection without saving
@app.post("/db-connections/test")
def test_db_connection(
    data: DBConnectionTest,
    current_user: User = Depends(get_current_user)
):
    try:
        conn_string = DatabaseConnector.build_connection_string(
            db_type=data.db_type,
            host=data.host,
            port=data.port,
            database=data.database,
            username=data.username,
            password=data.password,
            custom_string=data.connection_string
        )
        
        success, message = DatabaseConnector.test_connection(conn_string)
        
        return {
            "success": success,
            "message": message
        }
    except Exception as e:
        return {
            "success": False,
            "message": str(e)
        }

# Get tables from a database connection
@app.get("/db-connections/{conn_id}/tables")
def get_db_tables(
    conn_id: int,
    db=Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    from models import DatabaseConnection
    
    connection = db.query(DatabaseConnection).filter(
        DatabaseConnection.id == conn_id,
        DatabaseConnection.user_id == current_user.id
    ).first()
    
    if not connection:
        raise HTTPException(status_code=404, detail="Connection not found")
    
    try:
        tables = DatabaseConnector.get_tables(connection.connection_string)
        return {"tables": tables}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Get table info (schema)
@app.get("/db-connections/{conn_id}/tables/{table_name}/info")
def get_table_info(
    conn_id: int,
    table_name: str,
    db=Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    from models import DatabaseConnection
    
    connection = db.query(DatabaseConnection).filter(
        DatabaseConnection.id == conn_id,
        DatabaseConnection.user_id == current_user.id
    ).first()
    
    if not connection:
        raise HTTPException(status_code=404, detail="Connection not found")
    
    try:
        info = DatabaseConnector.get_table_info(connection.connection_string, table_name)
        return info
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Get sample data from table
@app.get("/db-connections/{conn_id}/tables/{table_name}/sample")
def get_table_sample(
    conn_id: int,
    table_name: str,
    limit: int = Query(default=10, ge=1, le=100),
    db=Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    from models import DatabaseConnection
    
    connection = db.query(DatabaseConnection).filter(
        DatabaseConnection.id == conn_id,
        DatabaseConnection.user_id == current_user.id
    ).first()
    
    if not connection:
        raise HTTPException(status_code=404, detail="Connection not found")
    
    try:
        sample = DatabaseConnector.get_sample_data(connection.connection_string, table_name, limit)
        
        # Update last_used timestamp
        connection.last_used = datetime.datetime.utcnow()
        db.commit()
        
        return sample
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Execute custom SQL query
class SQLQueryRequest(BaseModel):
    query: str
    limit: int | None = None

@app.post("/db-connections/{conn_id}/query")
def execute_sql_query(
    conn_id: int,
    data: SQLQueryRequest,
    db=Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    from models import DatabaseConnection
    
    connection = db.query(DatabaseConnection).filter(
        DatabaseConnection.id == conn_id,
        DatabaseConnection.user_id == current_user.id
    ).first()
    
    if not connection:
        raise HTTPException(status_code=404, detail="Connection not found")
    
    try:
        # Execute query
        df = DatabaseConnector.execute_query(
            connection.connection_string, 
            data.query, 
            limit=data.limit
        )
        
        # Update last_used timestamp
        connection.last_used = datetime.datetime.utcnow()
        db.commit()
        
        return {
            "columns": df.columns.tolist(),
            "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()},
            "data": df.to_dict('records'),
            "row_count": len(df)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Delete database connection
@app.delete("/db-connections/{conn_id}")
def delete_db_connection(
    conn_id: int,
    db=Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    from models import DatabaseConnection
    
    connection = db.query(DatabaseConnection).filter(
        DatabaseConnection.id == conn_id,
        DatabaseConnection.user_id == current_user.id
    ).first()
    
    if not connection:
        raise HTTPException(status_code=404, detail="Connection not found")
    
    db.delete(connection)
    db.commit()
    
    return {"msg": "Database connection deleted successfully"}

# Get supported database types
@app.get("/db-connections/supported-types")
def get_supported_db_types(current_user: User = Depends(get_current_user)):
    return {
        "types": DatabaseConnector.SUPPORTED_DB_TYPES,
        "details": {
            "mysql": {"default_port": 3306, "requires": ["pymysql"]},
            "postgresql": {"default_port": 5432, "requires": ["psycopg2-binary"]},
            "sqlite": {"default_port": None, "requires": []},
            "sqlserver": {"default_port": 1433, "requires": ["pyodbc"]},
            "oracle": {"default_port": 1521, "requires": ["cx_oracle"]}
        }
    }

# --- Chat History ---
from fastapi import Body

class ChatMessageRequest(PydanticBaseModel):
    session_id: str
    prompt: str
    response: str
    model_used: str | None = None  # Track which model generated the response

class ChatMessageResponse(PydanticBaseModel):
    id: int
    session_id: str
    prompt: str
    response: str
    timestamp: datetime.datetime
    model_used: str | None = None

from typing import List as TypingList

@app.post("/history", response_model=ChatMessageResponse)
def save_chat_message(
    data: ChatMessageRequest,
    db=Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    chat = ChatHistory(
        user_id=current_user.id,
        session_id=data.session_id,
        prompt=data.prompt,
        response=data.response,
        model_used=data.model_used  # Store which model was used
    )
    db.add(chat)
    db.commit()
    db.refresh(chat)
    return ChatMessageResponse(
        id=getattr(chat, 'id'),
        session_id=getattr(chat, 'session_id'),
        prompt=getattr(chat, 'prompt'),
        response=getattr(chat, 'response'),
        timestamp=getattr(chat, 'timestamp')
    )

from fastapi import Query as FastAPIQuery
from sqlalchemy import func

@app.get("/chat-sessions")
def list_chat_sessions(
    db=Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """List all chat sessions for the current user with their last message."""
    # Get distinct session_ids with their latest timestamp and message
    subquery = db.query(
        ChatHistory.session_id,
        func.max(ChatHistory.timestamp).label('latest_timestamp')
    ).filter(
        ChatHistory.user_id == current_user.id
    ).group_by(ChatHistory.session_id).subquery()
    
    # Join to get the full record for the latest message in each session
    sessions = db.query(ChatHistory).join(
        subquery,
        (ChatHistory.session_id == subquery.c.session_id) &
        (ChatHistory.timestamp == subquery.c.latest_timestamp)
    ).filter(
        ChatHistory.user_id == current_user.id
    ).order_by(ChatHistory.timestamp.desc()).all()
    
    return [
        {
            "session_id": chat.session_id,
            "last_message": chat.prompt or chat.response or "New Chat",
            "timestamp": chat.timestamp
        }
        for chat in sessions
    ]

@app.get("/history", response_model=TypingList[ChatMessageResponse])
def get_chat_history(
    session_id: str = FastAPIQuery(..., description="Session ID to filter chat history"),
    db=Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    chats = db.query(ChatHistory).filter(ChatHistory.user_id == current_user.id, ChatHistory.session_id == session_id).order_by(ChatHistory.timestamp.asc()).all()
    return [
        ChatMessageResponse(
            id=getattr(chat, 'id'),
            session_id=getattr(chat, 'session_id'),
            prompt=getattr(chat, 'prompt'),
            response=getattr(chat, 'response'),
            timestamp=getattr(chat, 'timestamp'),
            model_used=getattr(chat, 'model_used', None)
        ) for chat in chats
    ]

# Serve React build files (dist) as static files
if getattr(sys, 'frozen', False):
    # Running in a PyInstaller bundle
    frontend_build_dir = os.path.join(sys._MEIPASS, 'dist')
else:
    # Running in normal Python
    frontend_build_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'dist'))

print("Serving static files from:", frontend_build_dir)
print("index.html exists:", os.path.exists(os.path.join(frontend_build_dir, "index.html")))

from starlette.staticfiles import StaticFiles
from starlette.responses import FileResponse
from starlette.requests import Request

class SPAStaticFiles(StaticFiles):
    async def get_response(self, path: str, scope):
        print(f"Requested path: {path}")
        full_path = os.path.join(self.directory, path)
        if not os.path.exists(full_path) or os.path.isdir(full_path):
            index_path = os.path.join(self.directory, "index.html")
            print(f"Serving index.html for {path}")
            if os.path.exists(index_path):
                return FileResponse(index_path)
        response = await super().get_response(path, scope)
        print(f"Response status for {path}: {response.status_code}")
        return response

# --- Expiry Mechanism ---
EXPIRY_FILE = os.path.join(os.path.dirname(__file__), "install_time.txt")
EXPIRY_DAYS = 3

def check_expiry():
    if not os.path.exists(EXPIRY_FILE):
        with open(EXPIRY_FILE, "w") as f:
            f.write(str(time.time()))
        return False  # Not expired
    with open(EXPIRY_FILE, "r") as f:
        first_time = float(f.read().strip())
    expiry_time = first_time + EXPIRY_DAYS * 24 * 60 * 60
    if time.time() > expiry_time:
        return True  # Expired
    return False  # Not expired

# --- Per-user trial remaining endpoint ---
from fastapi import Depends
@app.get("/trial-remaining")
def trial_remaining(current_user: User = Depends(get_current_user)):
    TRIAL_DAYS = 3
    if not current_user.created_at:
        return {"remaining_seconds": 0}
    expiry_time = current_user.created_at + datetime.timedelta(days=TRIAL_DAYS)
    remaining = int((expiry_time - datetime.datetime.utcnow()).total_seconds())
    if remaining < 0:
        remaining = 0
    return {"remaining_seconds": remaining}

# Serve React build files (dist) as static files ONLY in development
# In Electron production mode, the main process serves the frontend directly
frontend_build_dir = None
if not os.environ.get("DESKTOP_MODE") == "1":
    # Only serve static files in development/server mode
    frontend_build_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'dist'))
    
    print("Serving static files from:", frontend_build_dir)
    print("index.html exists:", os.path.exists(os.path.join(frontend_build_dir, "index.html")))
    
    from starlette.staticfiles import StaticFiles
    from starlette.responses import FileResponse
    from starlette.requests import Request
    
    class SPAStaticFiles(StaticFiles):
        async def get_response(self, path: str, scope):
            print(f"Requested path: {path}")
            full_path = os.path.join(self.directory, path)
            if not os.path.exists(full_path) or os.path.isdir(full_path):
                index_path = os.path.join(self.directory, "index.html")
                print(f"Serving index.html for {path}")
                if os.path.exists(index_path):
                    return FileResponse(index_path)
            response = await super().get_response(path, scope)
            print(f"Response status for {path}: {response.status_code}")
            return response
    
    if os.path.exists(frontend_build_dir):
        app.mount("/", SPAStaticFiles(directory=frontend_build_dir, html=True), name="static")
    else:
        print(f"⚠️  Frontend build directory not found: {frontend_build_dir}")
else:
    print("Running in desktop mode - frontend served by Electron main process")

REQUIRED_MODELS = [
    # "llama3.1:latest",
    # "anindya/prem1b-sql-ollama-fp116:latest",
    # "deepseek-coder-v2:latest"  # Added for data analysis
    "gemma3:270m"
]

def check_model_exists(model_name):
    """Check if a model already exists locally"""
    try:
        result = subprocess.run(
            ["ollama", "list"], 
            capture_output=True, 
            text=True, 
            timeout=5
        )
        if result.returncode == 0:
            # Parse the output to see if model exists
            return model_name in result.stdout
        return False
    except Exception:
        return False

def ensure_models():
    """Only pull models if they don't already exist"""
    for model in REQUIRED_MODELS:
        try:
            if check_model_exists(model):
                print(f"✓ Model already available: {model}")
                continue
            
            print(f"Pulling model: {model}")
            subprocess.run(["ollama", "pull", model], check=True)
            print(f"✓ Successfully pulled: {model}")
        except Exception as e:
            print(f"Failed to pull model {model}: {e}")

# Only ensure models if not in desktop mode (to avoid blocking startup)
if os.environ.get("DESKTOP_MODE") != "1":
    # Ensure models are present before starting the server (in server mode)
    ensure_models()
else:
    print("Desktop mode: Skipping model pull at startup. Models will be pulled on demand.")

if __name__ == "__main__":
    import uvicorn
    # Use Render's PORT environment variable, fallback to 8000 for local development
    port = int(os.getenv("PORT", 8000))
    
    # Only open browser if not running from desktop launcher
    if not os.getenv("DESKTOP_MODE"):
        import threading
        import webbrowser
        def open_browser():
            webbrowser.open_new("http://localhost:8000")
        threading.Timer(1.5, open_browser).start()
    
    print(f"🚀 Starting server on port {port}")
    print(f"🌐 Environment: PORT={os.getenv('PORT', 'Not set (using 8000)')}")
    uvicorn.run(app, host="0.0.0.0", port=port)
